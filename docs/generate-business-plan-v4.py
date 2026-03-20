#!/usr/bin/env python3
"""
Generate SocialHomes.Ai Investment Business Plan v4 as a Word document.

Key changes from v3:
- GBP100M valuation by Year 6 (~GBP12.5M ARR at 8x)
- GBP260M valuation by Year 10 (GBP32.4M ARR)
- Fixed overlapping chart labels throughout
- AI engine = flexible (Claude + Vertex, future-proof)
- Self-service import/migration as key differentiator
- On-demand tenant provisioning (one-click GCP instances)
- Payment provider integration (GoCardless/Stripe)
- Embeddable tenant portal widgets
- TSM feedback collection built into tenant journey
- Public property data auto-enrichment
- 40% gross margin target from Year 4
- CTO throughout including minimum viable budget
- Product roadmap section with 8 new features
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

# Brand colours
TEAL = RGBColor(0, 170, 164)
DARK = RGBColor(30, 30, 30)
GREY = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)
LIGHT_GREY = RGBColor(240, 240, 240)

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
CHART_DIR = os.path.join(OUTPUT_DIR, 'charts')
LOGO_PATH = os.path.join(CHART_DIR, 'yantra-logo.png')
os.makedirs(CHART_DIR, exist_ok=True)

# ── Data constants ────────────────────────────────────────────────

YEARS_SHORT = ['Y1', 'Y2', 'Y3', 'Y4', 'Y5', 'Y6', 'Y7', 'Y8', 'Y9', 'Y10']

# Customer counts by segment
CUST_SMALL =      [10, 15, 22, 30, 38, 45, 50, 55, 58, 60]
CUST_MID =        [0,  0,  5, 14, 26, 40, 55, 68, 78, 88]
CUST_LARGE =      [0,  0,  0,  4, 10, 18, 26, 34, 42, 50]
CUST_ENTERPRISE = [0,  0,  0,  0,  3,  7, 12, 18, 24, 30]
CUST_G15 =        [0,  0,  0,  0,  0,  2,  4,  7, 10, 14]
CUST_TOTAL =      [10, 15, 27, 48, 77, 112, 147, 182, 212, 242]

# Revenue (GBP thousands) — REALISTIC: no SaaS Y1, minimal Y2, ramp Y3+
# Y1-2 = build years (consortium revenue only + 1-2 early SaaS Y2)
# Y3 = product launch, first real sales
# Y4+ = mid-market entry, scaling
ARR = [75, 125, 720, 2200, 5000, 12500, 15800, 21440, 26800, 32440]
REVENUE = ARR

# Costs: lean Y1-2 (founder unpaid Y1), 40% gross margin from Y4
COSTS = [109, 166, 353, 1320, 3000, 7500, 9480, 12860, 16080, 19460]
NET = [-34, -41, 367, 880, 2000, 5000, 6320, 8580, 10720, 12980]
HEADCOUNT = [3, 5, 10, 18, 28, 48, 60, 75, 88, 100]

# Valuations at 8x ARR (GBP millions)
VALUATIONS = [0.6, 1.0, 5.8, 17.6, 40, 100, 126.4, 171.5, 214.4, 259.5]

# ── Chart generation ──────────────────────────────────────────────

def create_market_segmentation_chart():
    """Pie chart with legend for small segments to avoid label overlap."""
    fig, (ax, ax_leg) = plt.subplots(1, 2, figsize=(8, 4.5),
                                      gridspec_kw={'width_ratios': [3, 2]})
    sizes = [1126, 180, 60, 50, 20]
    # Only label the two largest segments inline
    labels = ['', '', '', '', '']
    colors = ['#e0e0e0', '#00AAA4', '#7dd3d0', '#b0b0b0', '#005f5c']
    explode = (0, 0.08, 0.04, 0, 0)
    wedges, texts, autotexts = ax.pie(
        sizes, explode=explode, labels=labels, colors=colors,
        autopct='%1.0f%%', startangle=90, pctdistance=0.78,
        textprops={'fontsize': 8},
    )
    # Add inline labels only for the two big segments
    ax.annotate('< 1,000 homes\n(1,126)', xy=(0, -0.3), fontsize=7.5, ha='center')
    ax.annotate('1,000-5,000\n(180)', xy=(0.85, 0.6), fontsize=7, ha='center', color='#005f5c', fontweight='bold')
    for t in autotexts:
        t.set_fontsize(7)
        t.set_fontweight('bold')
    ax.set_title('UK Registered Providers\nby Stock Size', fontsize=11, fontweight='bold', pad=10)
    # Legend panel for all segments (avoids overlap)
    ax_leg.axis('off')
    legend_data = [
        ('#e0e0e0', '< 1,000 homes', '1,126 providers (78%)'),
        ('#00AAA4', '1,000\u20135,000', '180 providers (13%) \u2014 YEARS 1-2'),
        ('#7dd3d0', '5,000\u201315,000', '60 providers (4%) \u2014 YEARS 3-4'),
        ('#b0b0b0', '15,000\u201350,000', '50 providers (3%) \u2014 YEARS 5-7'),
        ('#005f5c', '50,000+', '20 providers (1%) \u2014 G15, YEARS 7+'),
    ]
    for i, (c, title, desc) in enumerate(legend_data):
        y = 0.88 - i * 0.17
        ax_leg.add_patch(plt.Rectangle((0.0, y - 0.03), 0.06, 0.06,
                         transform=ax_leg.transAxes, color=c))
        ax_leg.text(0.1, y + 0.01, title, transform=ax_leg.transAxes,
                    fontsize=8, fontweight='bold', va='center')
        ax_leg.text(0.1, y - 0.06, desc, transform=ax_leg.transAxes,
                    fontsize=7, va='center', color='#555')
    fig.text(0.5, 0.01, 'Source: RSH 2024-25. Total: 1,581 providers, 4.5M homes.',
             ha='center', fontsize=6.5, color='grey')
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_market_segmentation.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_arr_growth_chart():
    fig, ax = plt.subplots(figsize=(8, 5))
    x = np.arange(len(YEARS_SHORT))
    colors = ['#7dd3d0' if a < 1000 else '#00AAA4' if a < 5000 else '#005f5c' for a in ARR]
    bars = ax.bar(x, [a / 1000 for a in ARR], color=colors, width=0.65,
                  edgecolor='white', linewidth=0.5)
    for bar, a in zip(bars, ARR):
        label = f'\u00a3{a / 1000:.1f}M' if a >= 1000 else f'\u00a3{a}k'
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + bar.get_height() * 0.04 + 0.2,
                label, ha='center', fontsize=8, fontweight='bold', color='#1e1e1e')
    ax.set_ylabel('Annual Recurring Revenue (\u00a3M)', fontsize=10)
    ax.set_title('ARR Growth \u2014 \u00a3200k to \u00a332.4M', fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(YEARS_SHORT, fontsize=9)
    ax.set_ylim(0, max(ARR) / 1000 * 1.15)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'\u00a3{v:.0f}M'))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    # Key milestone annotations — offset to avoid overlap
    ax.annotate('\u00a3100M\nvaluation', xy=(5, ARR[5]/1000), xytext=(3.5, ARR[5]/1000 + 3),
                fontsize=8, fontweight='bold', color='#005f5c',
                arrowprops=dict(arrowstyle='->', color='#005f5c', lw=1.5))
    ax.annotate('\u00a3260M\nvaluation', xy=(9, ARR[9]/1000), xytext=(7.5, ARR[9]/1000 + 2),
                fontsize=8, fontweight='bold', color='#005f5c',
                arrowprops=dict(arrowstyle='->', color='#005f5c', lw=1.5))
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_arr_growth.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_customer_growth_stacked_chart():
    fig, ax = plt.subplots(figsize=(9, 5))
    x = np.arange(len(YEARS_SHORT))
    w = 0.55
    ax.bar(x, CUST_SMALL, w, label='Small (1-5k homes)', color='#7dd3d0')
    ax.bar(x, CUST_MID, w, bottom=CUST_SMALL, label='Mid-market (5-15k)', color='#00AAA4')
    bottom2 = [s + m for s, m in zip(CUST_SMALL, CUST_MID)]
    ax.bar(x, CUST_LARGE, w, bottom=bottom2, label='Large (15-30k)', color='#008a85')
    bottom3 = [b + l for b, l in zip(bottom2, CUST_LARGE)]
    ax.bar(x, CUST_ENTERPRISE, w, bottom=bottom3, label='Enterprise (30-50k)', color='#005f5c')
    bottom4 = [b + e for b, e in zip(bottom3, CUST_ENTERPRISE)]
    ax.bar(x, CUST_G15, w, bottom=bottom4, label='G15/Mega (50-125k)', color='#003d3b')
    for i, t in enumerate(CUST_TOTAL):
        ax.text(i, t + 3, str(t), ha='center', fontsize=8, fontweight='bold', color='#333333')
    ax.set_ylabel('Number of Customers', fontsize=10)
    ax.set_title('Customer Growth by Market Segment \u2014 10-Year', fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(YEARS_SHORT, fontsize=9)
    ax.legend(fontsize=8, loc='upper left')
    ax.set_ylim(0, 210)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_customer_growth_stacked.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_revenue_cost_3yr_chart():
    fig, ax = plt.subplots(figsize=(7, 4.5))
    years = ['FY 26/27\n(Year 1)', 'FY 27/28\n(Year 2)', 'FY 28/29\n(Year 3)']
    revenue = [200, 280, 720]
    costs = [375, 490, 750]
    x = np.arange(len(years))
    ax.plot(x, revenue, 'o-', color='#00AAA4', linewidth=2.5, markersize=8, label='Revenue')
    ax.plot(x, costs, 'o-', color='#cc3333', linewidth=2.5, markersize=8, label='Total Costs')
    ax.fill_between(x, revenue, costs,
                    where=[r < c for r, c in zip(revenue, costs)],
                    color='#ffcccc', alpha=0.3, label='Funding gap')
    ax.fill_between(x, revenue, costs,
                    where=[r >= c for r, c in zip(revenue, costs)],
                    color='#ccffcc', alpha=0.3, label='Surplus')
    for i, (r, c) in enumerate(zip(revenue, costs)):
        ax.annotate(f'\u00a3{r:,}k', (i, r), textcoords='offset points', xytext=(0, 12),
                    ha='center', fontsize=9, fontweight='bold', color='#00AAA4')
        ax.annotate(f'\u00a3{c:,}k', (i, c), textcoords='offset points', xytext=(0, -18),
                    ha='center', fontsize=9, fontweight='bold', color='#cc3333')
    ax.set_xticks(x)
    ax.set_xticklabels(years, fontsize=9)
    ax.set_ylabel('\u00a3 thousands', fontsize=9)
    ax.set_title('Revenue vs Costs \u2014 3-Year Detailed', fontsize=11, fontweight='bold')
    ax.legend(loc='upper left', fontsize=8)
    ax.set_ylim(0, 900)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'\u00a3{v:,.0f}k'))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.annotate('~Breakeven\nYear 3', xy=(2, 720), xytext=(1.3, 820),
                fontsize=8, fontweight='bold', color='#00AAA4',
                arrowprops=dict(arrowstyle='->', color='#00AAA4'))
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_revenue_cost_3yr.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_revenue_cost_10yr_chart():
    fig, ax = plt.subplots(figsize=(9, 5))
    x = np.arange(len(YEARS_SHORT))
    rev_m = [r / 1000 for r in REVENUE]
    cost_m = [c / 1000 for c in COSTS]
    ax.plot(x, rev_m, 'o-', color='#00AAA4', linewidth=2.5, markersize=7, label='Revenue')
    ax.plot(x, cost_m, 'o-', color='#cc3333', linewidth=2.5, markersize=7, label='Costs')
    ax.fill_between(x, rev_m, cost_m,
                    where=[r > c for r, c in zip(rev_m, cost_m)],
                    color='#ccffcc', alpha=0.3, label='Profit')
    ax.fill_between(x, rev_m, cost_m,
                    where=[r <= c for r, c in zip(rev_m, cost_m)],
                    color='#ffcccc', alpha=0.3, label='Loss')
    for i in range(len(x)):
        ax.text(i, rev_m[i] + 0.5, f'\u00a3{rev_m[i]:.1f}M', ha='center', fontsize=7,
                fontweight='bold', color='#00AAA4')
        ax.text(i, cost_m[i] - 0.8, f'\u00a3{cost_m[i]:.1f}M', ha='center', fontsize=7,
                fontweight='bold', color='#cc3333')
    ax.set_xticks(x)
    ax.set_xticklabels(YEARS_SHORT, fontsize=9)
    ax.set_ylabel('\u00a3 millions', fontsize=10)
    ax.set_title('Revenue vs Costs \u2014 10-Year Growth Model', fontsize=12, fontweight='bold')
    ax.legend(fontsize=8, loc='upper left')
    ax.set_ylim(0, 28)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'\u00a3{v:.0f}M'))
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_revenue_cost_10yr.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_pricing_comparison_chart():
    fig, ax = plt.subplots(figsize=(7, 4))
    vendors = [
        'Aareon QL\n(\u00a360k-\u00a3100k/yr)',
        'Civica Cx\n(\u00a350k-\u00a380k/yr)',
        'MRI / Orchard\n(\u00a340k-\u00a370k/yr)',
        'OmniLedger\n(\u00a320k-\u00a340k/yr)',
        'SocialHomes.Ai\n(\u00a320k/yr)',
    ]
    costs_mid = [80, 65, 55, 30, 20]
    colors = ['#808080', '#808080', '#808080', '#808080', '#00AAA4']
    bars = ax.barh(vendors, costs_mid, color=colors, height=0.6)
    ax.set_xlabel('Annual Cost (\u00a3000s) \u2014 midpoint estimate', fontsize=9)
    ax.set_title('Annual HMS Cost Comparison\n(For a ~2,500-home housing association)',
                 fontsize=11, fontweight='bold')
    for bar, cost in zip(bars, costs_mid):
        ax.text(bar.get_width() + 1.5, bar.get_y() + bar.get_height() / 2,
                f'~\u00a3{cost:,.0f}k', va='center', fontsize=9, fontweight='bold')
    ax.set_xlim(0, 110)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.text(0.5, -0.02, 'Sources: Sector benchmarks for providers with 1,000-5,000 homes.',
             ha='center', fontsize=6.5, color='grey')
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_pricing_comparison.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_market_entry_timeline_chart():
    fig, ax = plt.subplots(figsize=(10, 3.5))
    segments = [
        ('Small HAs\n1-5k homes', 0, 10, '#7dd3d0'),
        ('Mid-market\n5-15k homes', 2, 8, '#00AAA4'),
        ('Large\n15-30k homes', 4, 6, '#008a85'),
        ('Enterprise\n30-50k homes', 4, 6, '#005f5c'),
        ('G15 / Mega\n50-125k homes', 5, 5, '#003d3b'),
    ]
    for i, (label, start, duration, color) in enumerate(segments):
        ax.barh(i, duration, left=start, height=0.6, color=color,
                edgecolor='white', linewidth=0.5)
        ax.text(start + duration / 2, i, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color='white')
    ax.set_yticks([])
    ax.set_xticks(range(11))
    ax.set_xticklabels(
        ['Y1\n26/27', 'Y2', 'Y3', 'Y4', 'Y5', 'Y6', 'Y7', 'Y8', 'Y9', 'Y10', ''],
        fontsize=8)
    ax.set_xlabel('Year', fontsize=9)
    ax.set_title('Market Entry Timeline \u2014 Start Small, Scale Up to G15',
                 fontsize=12, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.set_xlim(-0.5, 10.5)
    plt.tight_layout()
    path = os.path.join(CHART_DIR, 'v4_market_entry_timeline.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


def create_tam_sam_som_chart():
    """TAM/SAM/SOM with legend outside to avoid overlapping labels."""
    fig, (ax, ax_legend) = plt.subplots(1, 2, figsize=(7.5, 4.5),
                                         gridspec_kw={'width_ratios': [3, 2]})
    # Draw circles with NO text inside
    for r, color, alpha in [(3.0, '#e0e0e0', 0.5), (2.0, '#7dd3d0', 0.6), (1.0, '#00AAA4', 0.7)]:
        ax.add_patch(plt.Circle((0, 0), r, color=color, alpha=alpha))
    # Minimal labels inside circles — just the tier name
    ax.text(0, 2.5, 'TAM', ha='center', va='center', fontsize=9, fontweight='bold', color='#555')
    ax.text(0, 1.3, 'SAM', ha='center', va='center', fontsize=9, fontweight='bold', color='#333')
    ax.text(0, 0, 'SOM', ha='center', va='center', fontsize=10, fontweight='bold', color='white')
    ax.set_xlim(-3.5, 3.5)
    ax.set_ylim(-3.5, 3.5)
    ax.set_aspect('equal')
    ax.axis('off')
    # Legend panel with full details — no overlap possible
    ax_legend.axis('off')
    legend_items = [
        ('#e0e0e0', 'TAM', '~1,500 organisations\n\u00a3150M\u2013\u00a3250M/yr'),
        ('#7dd3d0', 'SAM', '~400 organisations\n\u00a340M\u2013\u00a380M/yr\n(reachable in 5 years)'),
        ('#00AAA4', 'SOM (Year 3)', '27 customers\n\u00a31.14M ARR'),
    ]
    for i, (color, title, desc) in enumerate(legend_items):
        y = 0.82 - i * 0.3
        ax_legend.add_patch(plt.Rectangle((0.05, y - 0.06), 0.08, 0.08,
                            transform=ax_legend.transAxes, color=color))
        ax_legend.text(0.18, y, title, transform=ax_legend.transAxes,
                       fontsize=9, fontweight='bold', va='center')
        ax_legend.text(0.18, y - 0.1, desc, transform=ax_legend.transAxes,
                       fontsize=7.5, va='center', color='#555')
    fig.suptitle('Market Opportunity', fontsize=12, fontweight='bold', y=0.97)
    plt.tight_layout(rect=[0, 0.03, 1, 0.94])
    path = os.path.join(CHART_DIR, 'v4_tam_sam_som.png')
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path


# ── Document helpers ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='00AAA4'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, header_color)
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    return table


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 1:
        for run in heading.runs:
            run.font.color.rgb = TEAL
    return heading


def add_source_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = GREY
    run.italic = True


def add_bold_para(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


# ── Main document ─────────────────────────────────────────────────

def build_document():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════

    for _ in range(2):
        doc.add_paragraph()

    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.5))

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INVESTMENT BUSINESS PLAN')
    run.font.size = Pt(28)
    run.font.color.rgb = TEAL
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SocialHomes.Ai')
    run.font.size = Pt(20)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI-Native Housing Management for UK Social Housing')
    run.font.size = Pt(13)
    run.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('From Small HAs to the G15 \u2014 A 10-Year Growth Strategy')
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    run.italic = True

    for _ in range(2):
        doc.add_paragraph()

    details = [
        ('Prepared by:', 'Yantra Works Ltd'),
        ('Website:', 'yantra.works'),
        ('Date:', 'March 2026'),
        ('Version:', '3.0'),
        ('Classification:', 'COMMERCIALLY SENSITIVE & CONFIDENTIAL'),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label} ')
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, 'Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Market Overview: UK Social Housing',
        '3. The Problem: Why Providers Need a New Approach',
        '4. The Regulatory Burning Platform',
        '5. The Solution: SocialHomes.Ai',
        '6. What\'s Built vs What\'s Needed (Honest 50-55% Assessment)',
        '7. Product Roadmap & Differentiators',
        '8. AI Strategy: Flexible, Multi-Model Intelligence',
        '9. Competitive Landscape',
        '10. Business Model & Pricing (5 Tiers)',
        '11. Go-to-Market: From Small HAs to the G15',
        '12. Financial Projections \u2014 3-Year Detailed (40% Gross Margin from Y4)',
        '13. Financial Projections \u2014 10-Year Growth (\u00a3100M Valuation by Year 6)',
        '14. Team & Resourcing (Scales with Revenue)',
        '15. Minimum Viable Budget Scenario',
        '16. Risk Analysis',
        '17. Comparable Companies & Exit Scenarios',
        '18. Investment Ask (\u00a3250k for a \u00a3100M+ Opportunity)',
        '19. Key Milestones',
        'Appendix A: About Yantra Works Ltd',
        'Appendix B: Glossary for Non-Housing Readers',
        'Appendix C: Sources & References',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '1. Why SocialHomes.Ai?', level=1)

    # ── THE HOOK ──
    add_heading_styled(doc, 'The Problem in One Sentence', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        'A two-year-old boy called Awaab Ishak died because his housing association\'s systems '
        'were so poor that three years of mould complaints were lost in paperwork. '
        'SocialHomes.Ai exists to make sure that never happens again.'
    )
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        'The UK social housing sector manages 4.5 million homes for 9 million people. '
        'It is regulated by four new laws introduced since 2023, each demanding better data, '
        'faster response times, and real accountability. Yet 80% of the sector\'s technology was '
        'designed in the 1990s. Smaller housing associations \u2014 the ones closest to their communities \u2014 '
        'are priced out of modern systems, paying \u00a350,000\u2013\u00a3100,000 per year for software that '
        'wasn\'t built for them.'
    )

    add_heading_styled(doc, 'What Makes This Different', level=2)
    differences = [
        ('It\'s AI-native, not AI-bolted-on', 'Every existing housing system was built before AI existed. '
         'They bolt AI on as an expensive add-on module. SocialHomes.Ai weaves intelligence into every screen: '
         'it tells housing officers which tenants are at risk BEFORE they fall into crisis, '
         'which repairs will recur, and which complaints are heading to the Ombudsman.'),
        ('It\'s built for small providers, not enterprise leftovers', 'Civica, NEC, and MRI built their '
         'products for organisations managing 30,000+ homes. Smaller providers get the same software '
         'with features stripped out, at prices they can barely afford. SocialHomes.Ai is designed '
         'from the ground up for the 1,000\u20135,000 home segment \u2014 then scales upward.'),
        ('Self-service onboarding \u2014 no other vendor does this', 'Traditional HMS implementations take '
         '6\u201318 months and cost \u00a350,000\u2013\u00a3200,000. SocialHomes.Ai will offer guided self-service '
         'data import, automatic provisioning, and a go-live timeline measured in weeks, not years. '
         'This removes the single biggest barrier to switching.'),
        ('It combines public data with organisational data', 'The system automatically pulls weather forecasts, '
         'EPC ratings, flood risk, deprivation data, and crime statistics \u2014 and combines them with each '
         'organisation\'s own records. A housing officer opening a property record sees not just the tenancy '
         'and repairs history, but the damp risk based on tonight\'s weather, the area deprivation score, '
         'and whether the building\'s gas certificate is about to expire.'),
        ('It costs 50\u201370% less', 'For a 2,500-home housing association: \u00a320,000/year instead of '
         '\u00a355,000\u2013\u00a366,000. The savings go directly back into housing and services.'),
    ]
    for title, desc in differences:
        p = doc.add_paragraph()
        run = p.add_run(title + '. ')
        run.bold = True
        p.add_run(desc)

    add_heading_styled(doc, 'The Social Impact', level=2)
    doc.add_paragraph(
        'Better housing systems save lives. When a housing association can see that a vulnerable '
        'tenant hasn\'t been in contact for six months, it can send a welfare check before a crisis. '
        'When damp and mould complaints are tracked with legally mandated timers (Awaab\'s Law), '
        'children grow up in healthier homes. When AI predicts which properties need maintenance '
        'before they fail, emergency repairs drop and tenant satisfaction rises. '
        'This is not just a technology business \u2014 it is a direct intervention in the quality of life '
        'for millions of people in social housing.'
    )

    add_heading_styled(doc, 'The Business Opportunity', level=2)
    add_styled_table(doc, ['', ''], [
        ['Market', '1,581 housing providers managing 4.5M homes (\u00a327.4bn sector turnover)'],
        ['Target segment', '~400 providers with 1,000\u201330,000 homes \u2014 underserved by incumbents'],
        ['Regulatory driver', '4 new laws since 2023 (Awaab\'s Law, TSMs, Ombudsman Code, RSH inspections)'],
        ['Strategy', 'Start small (consortium of 10 HAs), prove product, expand to mid-market, then enterprise'],
        ['Year 3 revenue', '\u00a3720,000 (breakeven year \u2014 first mid-market sales)'],
        ['Year 6 target', '\u00a312.5M ARR \u2014 \u00a3100M valuation at 8x revenue'],
        ['Year 10 target', '\u00a332.4M ARR \u2014 242 customers across all segments'],
        ['Investment sought', '\u00a3250,000 \u2014 founder-led build with sales from Year 2. Profitable Year 3.'],
    ], col_widths=[4, 13])

    doc.add_paragraph()
    add_heading_styled(doc, 'Financial Summary', level=2)
    add_styled_table(doc, ['', 'FY 26/27', 'FY 27/28', 'FY 28/29'], [
        ['Revenue', '\u00a375,000', '\u00a3125,000', '\u00a3720,000'],
        ['Costs (incl. taxes & VAT)', '\u00a3235,000', '\u00a3352,000', '\u00a3557,000'],
        ['Net', '(\u00a3160,000)', '(\u00a3227,000)', '+\u00a3163,000'],
    ], col_widths=[5, 3.5, 3.5, 3.5])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'Investment sought: \u00a3250,000. The founder takes a \u00a395,000 salary (CTO role) from Year 1. '
        'All UK roles include employer NI (13.8%) and pension (3%). Offshore services include 20% VAT. '
        'Total funding: \u00a3250k investment + \u00a3150k consortium contributions = \u00a3400k. '
        'The business reaches profitability in Year 3 and becomes self-sustaining. '
        'No further funding rounds are needed unless accelerating growth into enterprise.'
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. MARKET OVERVIEW
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '2. Market Overview: UK Social Housing', level=1)

    doc.add_paragraph(
        'This section provides context for readers unfamiliar with the UK social housing sector.'
    ).runs[0].italic = True

    add_heading_styled(doc, 'What is Social Housing?', level=2)
    doc.add_paragraph(
        'Social housing is subsidised rental accommodation provided by two types of organisation in England:'
    )
    doc.add_paragraph(
        'Housing associations (also called Registered Providers or RPs) \u2014 not-for-profit organisations '
        'that own and manage social housing. There are 1,353 private registered providers in England.',
        style='List Bullet',
    )
    doc.add_paragraph(
        'Local authorities (councils) \u2014 228 local authorities are registered with the Regulator of Social Housing, '
        'of which approximately 143 still directly own and manage council housing stock.',
        style='List Bullet',
    )
    doc.add_paragraph(
        'Together, they manage approximately 4.5 million homes across England, housing roughly 9 million '
        'people \u2014 around 17% of all households. The sector has a combined annual turnover of \u00a327.4 billion.'
    )
    add_source_note(doc, 'Source: RSH Statistical Data Return 2024-25, RSH Global Accounts 2025')

    add_heading_styled(doc, 'Market Segmentation', level=2)
    doc.add_paragraph(
        'The sector is highly concentrated: the 19 largest providers manage 42% of all stock and generate '
        '47% of sector turnover. However, 83% of all private registered providers manage fewer than 1,000 homes. '
        'Our go-to-market strategy targets progressively larger segments over a 10-year period:'
    )

    chart_path = os.path.join(CHART_DIR, 'info_market_segmentation.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.2))
    else:
        chart_path = create_market_segmentation_chart()
        doc.add_picture(chart_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_table(doc, ['Segment', 'Providers', '% of Total', 'Our Entry Window'], [
        ['< 1,000 homes', '~1,126', '79%', 'Too small for full HMS'],
        ['1,000 \u2013 5,000 homes', '~180', '13%', 'YEARS 1\u20132: Primary target \u2014 prove the product'],
        ['5,000 \u2013 15,000 homes', '~60', '4%', 'YEARS 3\u20134: Mid-market \u2014 where the real revenue is'],
        ['15,000 \u2013 50,000 homes', '~50', '3%', 'YEARS 5\u20137: Large providers with bigger contracts'],
        ['50,000+ homes (G15)', '~20', '1%', 'YEARS 6\u201310: G15 members \u2014 \u00a3500k+/yr contracts'],
    ], col_widths=[4, 2.5, 2, 8.5])

    add_heading_styled(doc, 'The G15 Opportunity', level=2)
    doc.add_paragraph(
        'The G15 are London\'s 15 largest housing associations, collectively managing over 800,000 homes. '
        'They represent the pinnacle of the market and are the ultimate validation for any HMS vendor. '
        'Key G15 members include:'
    )
    add_styled_table(doc, ['G15 Member', 'Homes Managed', 'Annual Turnover'], [
        ['Clarion Housing Group', '~125,000', '\u00a3824M'],
        ['Peabody', '~108,000', '\u00a31.1B'],
        ['L&Q', '~105,000', '\u00a3995M'],
        ['Notting Hill Genesis', '~68,000', '\u00a3748M'],
        ['Hyde Group', '~44,000', '\u00a3382M'],
        ['Network Homes', '~20,000', '\u00a3206M'],
    ], col_widths=[5.5, 4, 4])
    doc.add_paragraph(
        'These organisations currently pay \u00a3500k\u2013\u00a32M+ per year for HMS solutions. Even capturing 10 G15/mega '
        'provider contracts at \u00a3500k/yr by Year 10 adds \u00a35M ARR \u2014 and validates the platform for the entire sector.'
    )
    add_source_note(doc, 'Source: G15 member annual reports, RSH Global Accounts 2025, Housing Today Largest 50 2025')

    add_heading_styled(doc, 'Local Authority Housing', level=2)
    doc.add_paragraph(
        'Approximately 143 local authorities in England still directly own and manage council housing, '
        'collectively owning over 1.4 million homes. Many smaller councils face the same challenges as '
        'housing associations. The number of ALMOs has fallen from 65 to just 19, as councils take management '
        'back in-house \u2014 increasing the need for cost-effective modern HMS solutions.'
    )
    add_source_note(doc, 'Source: ARCH Housing, RSH 2024-25, LGA')

    add_heading_styled(doc, 'Total Addressable Market', level=2)
    chart_path = os.path.join(CHART_DIR, 'info_tam_sam_som.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.8))
    else:
        chart_path = create_tam_sam_som_chart()
        doc.add_picture(chart_path, width=Inches(3.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc, ['Market Layer', 'Organisations', 'Annual Value'], [
        ['TAM (Total Addressable Market)', '~1,500 (all HAs + LAs)', '\u00a3150M \u2013 \u00a3250M/year'],
        ['SAM (Serviceable Addressable Market)', '~400 (reachable within 5 years)', '\u00a340M \u2013 \u00a380M/year'],
        ['SOM (Serviceable Obtainable Market, Year 3)', '24 customers', '\u00a3720k ARR'],
    ], col_widths=[6, 6, 5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. THE PROBLEM
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '3. The Problem: Why Providers Need a New Approach', level=1)

    doc.add_paragraph(
        'The regulatory, compliance, and operational demands on housing providers have never been greater. '
        'Yet most systems housing associations use to manage these obligations were designed in the 1990s and 2000s.'
    )

    add_heading_styled(doc, 'What Providers Currently Pay', level=2)
    doc.add_paragraph(
        'HMS costs vary dramatically by provider size. The table below shows typical costs across market segments:'
    )
    add_styled_table(doc, ['Provider Size', 'Typical Annual HMS Cost', 'Typical Implementation'], [
        ['Small (1,000\u20135,000 homes)', '\u00a340,000 \u2013 \u00a3100,000/yr', '\u00a315,000 \u2013 \u00a330,000'],
        ['Mid-market (5,000\u201315,000 homes)', '\u00a3100,000 \u2013 \u00a3250,000/yr', '\u00a340,000 \u2013 \u00a380,000'],
        ['Large (15,000\u201330,000 homes)', '\u00a3200,000 \u2013 \u00a3500,000/yr', '\u00a380,000 \u2013 \u00a3200,000'],
        ['Enterprise (30,000\u201350,000 homes)', '\u00a3400,000 \u2013 \u00a3800,000/yr', '\u00a3150,000 \u2013 \u00a3500,000'],
        ['G15/Mega (50,000\u2013125,000 homes)', '\u00a3500,000 \u2013 \u00a32,000,000+/yr', '\u00a3500,000 \u2013 \u00a32,000,000+'],
    ], col_widths=[5, 5.5, 5.5])
    add_source_note(doc, 'Source: Sector benchmarks, Find a Tender Service, procurement data, provider feedback.')

    doc.add_paragraph()
    doc.add_paragraph(
        'SocialHomes.Ai undercuts every segment by 50\u201375% \u2014 and delivers more AI capability than any incumbent.'
    )

    add_heading_styled(doc, 'The Gap in the Market', level=2)
    add_styled_table(doc, ['Current Option', 'Challenge'], [
        ['Enterprise HMS (Civica, NEC, MRI, Aareon)', 'Overpriced. Designed for 10,000+ home providers. AI bolted on as an afterthought. 30-year-old architectures.'],
        ['Budget HMS or spreadsheets', 'Cannot reliably meet Awaab\'s Law, TSM, Complaint Handling Code, or RSH Consumer Standards. Regulatory risk.'],
        ['Build your own', 'Too complex, expensive, and risky for any single organisation. No recent success stories.'],
        ['SocialHomes.Ai', 'Modern, AI-native, built for every segment. 5 pricing tiers. Consortium-governed.'],
    ], col_widths=[5, 12])

    add_heading_styled(doc, 'Legacy System Characteristics', level=2)
    doc.add_paragraph('The sector\'s dominant systems share common problems:')
    for problem in [
        'Client-server architecture designed before cloud computing \u2014 slow, forms-heavy, not mobile-friendly',
        'AI treated as an optional bolt-on module at additional cost, not integrated into workflows',
        'Proprietary APIs and restrictive commercial contracts that make integration difficult',
        'Data silos that contribute to regulatory failures \u2014 over 80% of regulatory downgrades cite data management issues',
        'NEC OHMS (Oracle-based) has announced end-of-life with de-support after March 2026, forcing migration',
    ]:
        doc.add_paragraph(problem, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. REGULATORY BURNING PLATFORM
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '4. The Regulatory Burning Platform', level=1)

    doc.add_paragraph(
        'Four major regulatory changes since 2023 are creating urgent, non-optional demand for better '
        'housing management systems. These are not aspirational improvements \u2014 they are legal requirements '
        'with enforcement powers including unlimited fines.'
    )

    add_heading_styled(doc, '4.1 Awaab\'s Law (From October 2025)', level=2)
    doc.add_paragraph(
        'Named after Awaab Ishak, a two-year-old who died in December 2020 from a severe respiratory '
        'condition caused by prolonged mould exposure in his Rochdale social housing home. His family had '
        'complained for three years without adequate response.'
    )
    doc.add_paragraph('Awaab\'s Law introduces legally binding response timescales:')
    add_styled_table(doc, ['Requirement', 'Timescale'], [
        ['Investigation of reported hazards', '10 working days'],
        ['Written summary to tenant after investigation', '3 working days'],
        ['Emergency hazard response', '24 hours'],
        ['Begin safety repair work', '5 working days after investigation'],
        ['Long-stop for supplementary/preventative work', '12 weeks'],
    ], col_widths=[10, 5])
    doc.add_paragraph()
    doc.add_paragraph(
        'Phased rollout: Phase 1 (October 2025) covers damp, mould, and emergency hazards. '
        'Phase 2 (2026) adds excess cold/heat, falls, fire, and electrical hazards. '
        'Phase 3 (2027) covers all remaining HHSRS hazards. '
        'The RSH can impose unlimited fines (cap removed by the 2023 Act), issue enforcement notices, '
        'and authorise emergency remedial action at the landlord\'s expense.'
    )
    add_source_note(doc, 'Source: GOV.UK Awaab\'s Law Guidance, Social Housing (Regulation) Act 2023')

    add_heading_styled(doc, '4.2 Tenant Satisfaction Measures (From April 2023)', level=2)
    doc.add_paragraph(
        '22 mandatory measures (12 perception-based, 10 management data) that all providers with 1,000+ homes '
        'must collect and report annually to the RSH. The 2024-25 headline results from 360 large landlords show:'
    )
    add_styled_table(doc, ['TSM Measure', '2024-25 Result'], [
        ['Overall tenant satisfaction (rental)', '71.8% median'],
        ['Home safety', '77.6%'],
        ['Fair treatment', '77.9%'],
        ['Complaint handling satisfaction', '35.5% (lowest score)'],
        ['Gas safety compliance', '99.7%'],
        ['Non-emergency repairs within target', '79.0%'],
    ], col_widths=[10, 5])
    add_source_note(doc, 'Source: RSH TSM 2024-25 Headline Report')

    add_heading_styled(doc, '4.3 Housing Ombudsman Complaint Handling Code (From April 2024)', level=2)
    doc.add_paragraph(
        'Became statutory on 1 April 2024. All landlords must follow a two-stage complaints process with '
        'strict timescales (10 working days for Stage 1, 20 for Stage 2). In 2024-25, the Housing Ombudsman '
        'made 7,082 determinations (up 30% year-on-year), ordered \u00a35.4 million in compensation, and issued '
        '168 Complaint Handling Failure Orders.'
    )
    add_source_note(doc, 'Source: Housing Ombudsman Annual Complaints Review 2024-25')

    add_heading_styled(doc, '4.4 RSH Consumer Standards (From April 2024)', level=2)
    doc.add_paragraph(
        'The Regulator of Social Housing began proactive consumer regulation inspections for all landlords '
        'with 1,000+ homes on a 4-year cycle. In the first year, at least 15 landlords received C3 gradings '
        '(serious failings). The RSH received 986 referrals in 2023-24, with 217 progressing to formal investigation.'
    )
    add_source_note(doc, 'Source: RSH Consumer Regulation Review 2023-24, GOV.UK')

    doc.add_paragraph()
    add_bold_para(doc,
        'Key insight for investors: These regulations are not optional. Housing associations that cannot '
        'demonstrate compliance risk unlimited fines, regulatory intervention, and public censure. '
        'This creates a compulsory purchasing trigger \u2014 organisations must upgrade their systems regardless '
        'of budget constraints.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. THE SOLUTION
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '5. The Solution: SocialHomes.Ai', level=1)

    doc.add_paragraph(
        'SocialHomes.Ai is an AI-native housing management system designed from the ground up for the '
        'entire UK social housing sector \u2014 from small housing associations to the G15. It combines modern '
        'cloud architecture with Anthropic\'s Claude AI to deliver intelligent, regulation-aware housing '
        'management at a fraction of incumbent costs.'
    )

    add_heading_styled(doc, 'Technology Stack', level=2)
    add_styled_table(doc, ['Layer', 'Technology'], [
        ['Frontend', 'React 19, TypeScript, Tailwind CSS 4'],
        ['Backend', 'Node.js, Express, TypeScript'],
        ['Database', 'Google Cloud Firestore (NoSQL, real-time sync)'],
        ['Authentication', 'Firebase Authentication (Google SSO + email/password)'],
        ['Hosting', 'Google Cloud Run (London, europe-west2)'],
        ['Data Standard', 'HACT UK Housing Data Standard v3.5'],
        ['AI Engine', 'Anthropic Claude (Sonnet 4 + Haiku 4.5) via API'],
        ['Data Residency', 'UK only (London region)'],
        ['Security', 'Google Cloud Armor WAF, TLS 1.3, AES-256 encryption at rest'],
        ['Compliance', 'SOC 2 Type II (Google Cloud), ISO 27001 infrastructure, GDPR'],
    ], col_widths=[4.5, 12])

    add_heading_styled(doc, 'Modules Designed', level=2)
    doc.add_paragraph(
        'The following modules have been designed, with full UI/UX across 40+ screens. '
        'See Section 6 for an honest assessment of what is working versus what remains to be built.'
    )
    add_styled_table(doc, ['Module', 'Screens', 'Key Features'], [
        ['Dashboard', '1', '6 KPI cards, interactive charts, activity feed, persona-scoped'],
        ['Daily Briefing', '1', 'Personalised by role, weather-linked alerts, urgent cases'],
        ['Tenancies', '3+', 'Searchable lists, UC tracking, vulnerability flags, AI risk profiles'],
        ['Properties', '3+', 'Map view, EPC ratings, damp risk scoring, compliance RAG'],
        ['Repairs', '3+', 'SLA tracking, priority triage, AI first-time-fix estimation'],
        ['Compliance', '3+', 'Big 6 safety, Awaab\'s Law countdown timers, audit trail'],
        ['Rent & Income', '2+', 'AI-prioritised arrears worklist, payment tracking'],
        ['Complaints', '2+', 'Housing Ombudsman Code compliant, stage tracking, TSM metrics'],
        ['ASB Cases', '2+', 'Case management with full timelines'],
        ['Communications', '2+', 'GOV.UK Notify templates, AI analysis'],
        ['Reports', '3+', '30+ reports including TSM, CORE, RSH SDR, Board Packs'],
        ['AI Centre', '2+', 'Insights, predictions, conversational AI assistant'],
        ['Admin', '10+', 'Users, teams, GDPR, audit, workflows'],
        ['Tenant Portal', '5+', 'Self-service with action buttons'],
        ['Explore / Map', '1', 'Interactive Leaflet map with property markers'],
        ['Allocations', '2+', 'Voids and lettings management'],
    ], col_widths=[3.5, 1.5, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. WHAT'S BUILT VS WHAT'S NEEDED (50-55%)
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '6. What\'s Built vs What\'s Needed', level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        'Honest assessment: The prototype is 50\u201355% complete \u2014 substantially further along than '
        'a typical early-stage startup. The codebase has 157 API route endpoints across 23 route files, '
        'with backend WRITE operations already working for tenants, properties, cases, admin, lettings, '
        'and booking (real Firestore writes, not stubs).'
    )
    run.bold = True
    run.italic = True

    add_heading_styled(doc, 'What\'s Working (50\u201355% of Production Scope)', level=2)

    working_items = [
        ('157 API route endpoints across 23 route files',
         'A comprehensive REST API covering all major housing management domains. This is not scaffolding \u2014 '
         'these are implemented endpoints with validation, error handling, and Firestore integration.'),
        ('Backend WRITE operations already exist',
         'Real Firestore writes for tenants, properties, cases, admin, lettings, and booking. '
         'POST and PATCH methods are implemented and tested.'),
        ('Frontend API client with POST/PATCH methods defined',
         'The API client layer has methods defined for creating and updating records. '
         'The gap is mainly that some frontend forms don\'t call these existing methods \u2014 '
         'this is pattern-following work that AI agents excel at.'),
        ('Full UI/UX across 40+ screens',
         'Complete responsive design system with dark theme, Tailwind CSS 4, '
         'persona-scoped layouts for 5 roles. Every screen is navigable.'),
        ('Awaab\'s Law countdown timer logic',
         'Real calculations from repair data \u2014 not mockups. '
         'The countdown timers compute remaining working days against statutory timescales.'),
        ('Authentication framework',
         'Firebase Auth with JWT tokens, RBAC roles (Housing Officer, Manager, Director, Admin, Tenant). '
         'Login/logout works. Role-based route guards in place.'),
        ('625 automated tests (architecture verified)',
         '228 server tests (Vitest), 219 E2E tests, 203 accessibility tests. All passing.'),
        ('CI/CD pipeline',
         'GitHub Actions triggers Cloud Build, deploys to Cloud Run (europe-west2). Automated.'),
        ('HACT v3.5 data schema defined',
         'Firestore collections follow the UK Housing Data Standard. Clear target schema for data migration.'),
    ]
    for title, detail in working_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(detail)
        run.font.size = Pt(9)

    add_heading_styled(doc, 'What AI Agents Can Build (80% of Remaining Work)', level=2)
    doc.add_paragraph(
        'The Yantra Framework\'s 7-agent pipeline is specifically designed for the pattern-following '
        'work that makes up the bulk of the remaining development:'
    )
    for item in [
        'Wire frontend forms to existing API endpoints \u2014 the API methods exist, forms just need connecting',
        'Add workflow state transitions \u2014 repair status progressions, complaint stage tracking, ASB case flow',
        'Implement multi-tenancy \u2014 org-scoping on Firestore queries (well-documented pattern)',
        'Connect Claude API for AI features \u2014 predictions, drafting, analysis, daily briefing',
        'Build remaining 28 reports \u2014 TSM, CORE, SDR, Board Packs with real data aggregation',
        'Complete admin CRUD \u2014 user management, team assignment, workflow configuration',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, 'What Humans Must Do (20% of Remaining Work)', level=2)
    doc.add_paragraph(
        'Some work requires human expertise that cannot be delegated to AI agents:'
    )
    for item in [
        'Independent QA testing company \u2014 functional testing, regression testing, UAT',
        'CREST penetration testing \u2014 security testing by certified ethical hackers',
        'Code review of AI-generated output \u2014 human oversight of every production commit',
        'Data migration from legacy HMS \u2014 understanding Civica/NEC/Aareon data models',
        'ISO 27001 implementation \u2014 information security management system',
        'Customer onboarding and training \u2014 HA-specific configuration, staff training',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_bold_para(doc,
        'Why this matters: At 50\u201355% complete with real backend writes working, '
        'the remaining development is de-risked. An investor is funding the completion of a working system '
        '\u2014 not a speculative idea. The build risk is substantially lower than a typical seed investment.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. PRODUCT ROADMAP & DIFFERENTIATORS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '7. Product Roadmap & Differentiators', level=1)

    doc.add_paragraph(
        'Beyond the core HMS modules, SocialHomes.Ai will deliver features that no incumbent offers. '
        'These are designed to dramatically lower barriers to adoption and create deep product moats.'
    )

    add_heading_styled(doc, 'a) Self-Service Onboarding & Data Import (Year 2)', level=2)
    doc.add_paragraph(
        'A guided self-service wizard for data migration from CSV, Excel, or direct API connection to legacy systems. '
        'Auto-mapping of fields to the HACT v3.5 standard with data validation, cleansing, and deduplication. '
        'No other HMS vendor offers fully hands-off self-service migration. This dramatically reduces '
        'implementation cost (from \u00a315\u201375k to near-zero for small organisations) and enables rapid '
        'customer onboarding at scale.'
    )

    add_heading_styled(doc, 'b) One-Click Tenant Provisioning (Year 2)', level=2)
    doc.add_paragraph(
        'Automated GCP project and Firestore instance creation per organisation. Each organisation gets a fully '
        'isolated database (not just collection scoping). Sign up \u2192 configure \u2192 import data \u2192 go live '
        'in days, not months. Infrastructure-as-code (Terraform/Pulumi) for repeatable provisioning. '
        'This is a significant competitive moat \u2014 incumbents take 6\u201318 months for implementation.'
    )

    add_heading_styled(doc, 'c) Payment Integration (Year 2-3)', level=2)
    doc.add_paragraph(
        'Connect to GoCardless for Direct Debit mandate collection, Stripe for card payments, and '
        'Open Banking APIs for real-time balance verification. Universal Credit managed payment tracking '
        'via DWP API. Payment reconciliation dashboard with automated matching.'
    )

    add_heading_styled(doc, 'd) Embeddable Tenant Widgets (Year 2-3)', level=2)
    doc.add_paragraph(
        'JavaScript widgets that organisations can embed on their own websites and portals: repair reporting, '
        'rent balance/payment, satisfaction surveys (feeding TSM data), communication preferences. '
        'These extend reach without requiring tenants to log into SocialHomes.Ai directly.'
    )

    add_heading_styled(doc, 'e) TSM Feedback Collection (Year 2)', level=2)
    doc.add_paragraph(
        'Built into the tenant journey at key touchpoints (post-repair, post-complaint, periodic). '
        'Multi-channel collection via SMS, email, widget, and in-app. Auto-feeds the 22 TSM measures '
        'into regulatory reports. Feeds individual tenant records with sentiment tracking over time. '
        'Benchmarking against published RSH sector averages.'
    )

    add_heading_styled(doc, 'f) Public Data Auto-Enrichment (Year 1-2)', level=2)
    doc.add_paragraph(
        'Automatically pull and combine publicly available data with each organisation\u2019s own data '
        'for richer insights. Sources include: EPC ratings from the Open Data register, Gas Safety and '
        'Electrical Safety Register lookups, Land Registry valuations, Environment Agency flood risk data, '
        'ONS Census 2021 and IMD/deprivation data, Police.uk crime data, Met Office weather forecasts '
        'for damp risk prediction. This creates intelligence that no single data source could provide alone.'
    )

    add_heading_styled(doc, 'g) Peer Comparison & Benchmarking (Year 3)', level=2)
    doc.add_paragraph(
        'Using RSH Statutory Data Returns (publicly available), organisations can compare their performance '
        'against sector averages: cost per unit, TSM scores, rent collection rates, repairs response times, '
        'void turnaround. Anonymous cross-consortium benchmarking for SocialHomes.Ai customers.'
    )

    add_heading_styled(doc, 'h) Predictive Maintenance (Year 3-4)', level=2)
    doc.add_paragraph(
        'AI analyses component age, repair history, building type, and weather exposure to predict when '
        'capital works are needed (roof, boiler, windows, etc.). Shifts from reactive repairs to planned '
        'maintenance, reducing emergency repair costs. Integrates with asset management and planned works scheduling.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. AI STRATEGY: FLEXIBLE, MULTI-MODEL INTELLIGENCE
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '8. AI Strategy: Flexible, Multi-Model Intelligence', level=1)

    doc.add_paragraph(
        'SocialHomes.Ai uses AI across three distinct layers. The AI engine is model-agnostic and '
        'future-proof \u2014 currently leveraging Anthropic Claude and Google Vertex AI, switching '
        'according to need, cost, and advances in the field.'
    )

    add_heading_styled(doc, 'Layer 1: AI BUILDS the Product', level=2)
    doc.add_paragraph(
        'The Yantra Framework is a proprietary 7-agent pipeline where AI agents '
        '(Business Analyst, Fullstack Developer, Frontend Developer, UI/UX Designer, QA Tester, '
        'Cyber Security Analyst, DevOps Engineer) collaborate under human supervision. This hybrid '
        'workforce approach has delivered the prototype at approximately 40% of the cost of traditional '
        'development. An independent testing company provides human verification \u2014 ensuring that '
        'production code is not solely reviewed by the same AI that wrote it.'
    )

    add_heading_styled(doc, 'Layer 2: AI POWERS the Product', level=2)
    doc.add_paragraph(
        'Claude is the application\'s AI engine, accessed via API. Every intelligent feature in '
        'SocialHomes.Ai is powered by Claude:'
    )
    for item in [
        'Predictive arrears \u2014 analyses payment patterns, UC status, and tenancy data to identify at-risk tenancies',
        'Repair prediction \u2014 analyses building age, component lifecycles, and repair history to predict failures',
        'Case management drafting \u2014 generates investigation plans, response letters, and escalation recommendations',
        'Report generation \u2014 compiles TSM, CORE, and SDR reports from raw data with narrative summaries',
        'Tenant risk profiling \u2014 identifies vulnerability indicators from interaction patterns',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_heading_styled(doc, 'Layer 3: AI AUGMENTS the User', level=2)
    doc.add_paragraph(
        'AI doesn\'t replace housing officers \u2014 it makes them dramatically more effective:'
    )
    for title, detail in [
        ('AI surfaces what matters', 'Daily briefing identifies the 5\u201310 things that need attention today, '
         'ranked by urgency and risk. Officers no longer need to search through queues.'),
        ('AI pre-fills and prepares', 'When an officer opens a case, AI has already drafted the investigation '
         'plan, pre-filled the response template, and suggested the timeline.'),
        ('AI drafts communications', 'Claude drafts tenant letters, complaint responses, and internal reports. '
         'Estimated time saving: 30\u201340% on administrative tasks.'),
        ('AI identifies hidden patterns', 'Recurring repair patterns suggesting capital works, silent tenants '
         'who may be in distress, arrears trajectories before they become critical.'),
        ('Future: supervised autonomous actions', 'With explicit org opt-in, selected safe actions '
         '(e.g., scheduling routine inspections, sending confirmations) can be automated \u2014 '
         'always logged, always auditable, always reversible.'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(detail)
        run.font.size = Pt(9)

    add_heading_styled(doc, 'Claude API Costs \u2014 Negligible at Scale', level=2)
    doc.add_paragraph('AI costs are a rounding error in the overall cost structure:')
    add_styled_table(doc, ['Scale', 'Users', 'Queries/Day', 'Monthly Claude Cost'], [
        ['Year 1 (10 small HAs)', '~200', '~500', '~\u00a345'],
        ['Year 3 (24 HAs)', '~600', '~1,500', '~\u00a395'],
        ['Year 5 (60 HAs)', '~2,000', '~5,000', '~\u00a3165'],
        ['Year 10 (185 HAs)', '~10,000', '~25,000', '~\u00a3256'],
    ], col_widths=[5, 2.5, 3, 4])
    add_source_note(doc, 'Based on Claude Sonnet 4: $3/MTok in, $15/MTok out. Haiku 4.5: $1/MTok in, $5/MTok out. 80% Haiku, 20% Sonnet mix.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. COMPETITIVE LANDSCAPE
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '8. Competitive Landscape', level=1)

    doc.add_paragraph(
        'This analysis covers the full market spectrum \u2014 from small HAs to the G15 \u2014 reflecting '
        'our 10-year growth ambition.'
    )

    add_heading_styled(doc, 'Incumbent Vendors by Segment', level=2)
    add_styled_table(doc, ['Vendor', 'Product', 'Primary Segment', 'Typical Annual Cost', 'AI'], [
        ['Aareon', 'QL', 'Small-Mid (250+ UK)', '\u00a360k\u2013\u00a3100k (small)', 'Limited'],
        ['Civica', 'Cx Housing', 'Mid-Large', '\u00a380k\u2013\u00a3250k', 'Bolt-on'],
        ['MRI / Orchard', 'ActiveH', 'Mid-Large', '\u00a370k\u2013\u00a3300k', 'Bolt-on'],
        ['NEC', 'OHMS (end-of-life)', 'Large-Enterprise', '\u00a3200k\u2013\u00a3800k', 'None'],
        ['Capita / IBS', 'OpenHousing', 'Enterprise-G15', '\u00a3300k\u2013\u00a31.5M', 'Minimal'],
        ['OmniLedger', 'PyramidG2', 'Small (budget)', '\u00a320k\u2013\u00a340k', 'None'],
        ['SocialHomes.Ai', 'Full platform', 'All segments', '\u00a320k\u2013\u00a3500k', 'Built-in (Claude)'],
    ], col_widths=[2.5, 2.5, 3, 3.5, 4])

    doc.add_paragraph()
    chart_path = os.path.join(CHART_DIR, 'info_pricing_comparison.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
    else:
        chart_path = create_pricing_comparison_chart()
        doc.add_picture(chart_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_styled(doc, 'Feature Comparison', level=2)
    add_styled_table(doc, ['Capability', 'Aareon QL', 'Civica Cx', 'MRI', 'SocialHomes.Ai'], [
        ['Cloud-native', 'Hosted', 'Partial', 'Partial', 'Yes (GCP)'],
        ['AI integrated into workflows', 'No', 'Bolt-on', 'Bolt-on', 'Yes (Claude)'],
        ['Awaab\'s Law countdown timers', 'Patch', 'Patch', 'Patch', 'Built-in'],
        ['TSM reporting', 'Module', 'Module', 'Module', 'Built-in'],
        ['Peer benchmarking (RSH data)', 'No', 'No', 'No', 'Yes'],
        ['Predictive maintenance AI', 'No', 'No', 'No', 'Yes'],
        ['HACT v3.5 native data model', 'No', 'No', 'No', 'Yes'],
        ['Daily AI briefing', 'No', 'No', 'No', 'Yes'],
        ['Annual cost (~2,500 homes)', '\u00a360\u2013100k', '\u00a350\u201380k', '\u00a340\u201370k', '\u00a320k'],
    ], col_widths=[5, 2.5, 2.5, 2.5, 4.5])

    add_heading_styled(doc, 'Defensible Moats', level=2)
    for moat in [
        'AI-native architecture \u2014 Cannot be replicated by bolting AI onto 30-year-old codebases',
        'HACT v3.5 from the data model up \u2014 First HMS built on the UK Housing Data Standard natively',
        'Consortium governance \u2014 Customers shape the roadmap via a Board seat',
        'Network effects \u2014 More organisations = better AI models via anonymised benchmarking',
        'Cost structure \u2014 AI-assisted development + offshore team delivers significantly lower build cost',
        'Regulatory head-start \u2014 Awaab\'s Law, TSM, Complaint Code built in from the start',
        'Full-spectrum pricing \u2014 5 tiers from \u00a320k to \u00a3500k covers every provider size',
    ]:
        doc.add_paragraph(moat, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. BUSINESS MODEL & PRICING (5 TIERS)
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '9. Business Model & Pricing', level=1)

    add_heading_styled(doc, 'Five-Tier Pricing Model', level=2)
    doc.add_paragraph(
        'Pricing scales with provider size, reflecting the complexity of implementation, the number '
        'of users, and the value delivered. Each tier undercuts incumbent vendors by 50\u201375%.'
    )
    add_styled_table(doc, ['Segment', 'Homes', 'Annual SaaS', 'Implementation', 'Price/Home/Yr'], [
        ['Small', '1,000 \u2013 5,000', '\u00a320,000', '\u00a315,000', '\u00a34 \u2013 \u00a320'],
        ['Mid-market', '5,000 \u2013 15,000', '\u00a380,000', '\u00a340,000', '\u00a35.33 \u2013 \u00a316'],
        ['Large', '15,000 \u2013 30,000', '\u00a3150,000', '\u00a375,000', '\u00a35 \u2013 \u00a310'],
        ['Enterprise', '30,000 \u2013 50,000', '\u00a3250,000', '\u00a3120,000', '\u00a35 \u2013 \u00a38.33'],
        ['G15/Mega', '50,000 \u2013 125,000', '\u00a3500,000', '\u00a3200,000', '\u00a34 \u2013 \u00a310'],
    ], col_widths=[2.5, 3, 3, 3, 3.5])

    doc.add_paragraph()
    doc.add_paragraph(
        'This pricing structure means our smallest customers pay \u00a320k/yr (vs \u00a340\u2013100k for incumbents) '
        'while our largest G15 customers pay \u00a3500k/yr (vs \u00a31\u20132M+ for NEC/Capita). The value proposition '
        'is consistent across segments: better technology at a lower price.'
    )

    add_heading_styled(doc, 'Revenue Streams', level=2)

    doc.add_paragraph('Stream 1: Consortium Co-Development Contributions', style='List Bullet')
    doc.add_paragraph(
        '10 founding housing associations each contribute \u00a315,000 over 24 months in quarterly '
        'instalments of \u00a31,875. Total: \u00a3150,000. This represents less than 0.1% of a typical '
        'small HA\'s annual revenue. In return: perpetual licence, code access, Board seat, '
        'influence over roadmap priorities.'
    )

    doc.add_paragraph('Stream 2: SaaS Subscriptions (5 tiers as above)', style='List Bullet')
    doc.add_paragraph(
        'Annual subscriptions ranging from \u00a320,000 for small HAs to \u00a3500,000 for G15 members. '
        'All-inclusive \u2014 AI features, hosting, updates, and support are included in the subscription.'
    )

    doc.add_paragraph('Stream 3: Implementation & Migration Fees (one-off)', style='List Bullet')
    doc.add_paragraph(
        'Data migration, system configuration, and staff training. Ranges from \u00a315,000 (small) '
        'to \u00a3200,000 (G15). These are one-off revenues that provide early cash flow.'
    )

    doc.add_paragraph('Stream 4: Managed Hosting & Support (annual)', style='List Bullet')
    doc.add_paragraph(
        'Premium support tiers, dedicated account management for enterprise customers, '
        'SLA-backed response times. Included in the SaaS subscription for standard support.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. GO-TO-MARKET: FROM SMALL HAs TO G15
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '10. Go-to-Market: From Small HAs to the G15', level=1)

    doc.add_paragraph(
        'The market entry strategy is deliberately staged. Each phase builds credibility and '
        'reference customers that unlock the next market segment.'
    )

    chart_path = os.path.join(CHART_DIR, 'info_market_entry.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
    else:
        chart_path = create_market_entry_timeline_chart()
        doc.add_picture(chart_path, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_styled(doc, 'Years 1\u20132: Small HAs (1,000\u20135,000 homes) \u2014 Prove the Product', level=2)
    doc.add_paragraph(
        '10 founding consortium members form the initial customer base. These are organisations '
        'large enough to need a proper HMS but small enough to take a risk on a new provider. '
        'The \u00a315,000 contribution (quarterly instalments of \u00a31,875) lowers the barrier to commitment. '
        'Channel: Direct outreach via 3C (sector engagement partner) with existing relationships '
        'across 50+ small-to-mid housing associations.'
    )

    add_heading_styled(doc, 'Years 3\u20134: Enter Mid-Market (5,000\u201315,000 homes)', level=2)
    doc.add_paragraph(
        'With 14\u201320 small HA reference customers and a production-proven platform, we enter the '
        'mid-market segment. These organisations pay \u00a380k/yr \u2014 4x the small HA tier. This is where '
        'the revenue model inflects. By end of Year 4: 40 total customers, \u00a31.66M ARR.'
    )

    add_heading_styled(doc, 'Years 5\u20137: Large Providers and G15 Entry (15,000\u2013125,000 homes)', level=2)
    doc.add_paragraph(
        'ISO 27001 certification, G-Cloud listing, and 40+ reference customers make SocialHomes.Ai '
        'eligible for large provider procurement. The first G15 member signs in Year 6 at \u00a3500k/yr. '
        'By end of Year 7: 109 total customers, \u00a39.5M ARR.'
    )

    add_heading_styled(doc, 'Years 8\u201310: Market Dominance', level=2)
    doc.add_paragraph(
        'The platform is proven across all segments. Network effects from anonymised benchmarking '
        'data make SocialHomes.Ai increasingly valuable to each new customer. '
        'By Year 10: 242 customers, \u00a332.4M ARR, 10 G15/mega provider contracts.'
    )

    add_heading_styled(doc, '10-Year Customer Growth by Segment', level=2)
    chart_path = os.path.join(CHART_DIR, 'info_customer_growth.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
    else:
        chart_path = create_customer_growth_stacked_chart()
        doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_table(doc, ['Year', 'Small', 'Mid', 'Large', 'Ent.', 'G15', 'Total', 'ARR'], [
        ['Y1 (FY26/27)', '10', '0', '0', '0', '0', '10', '\u00a3200k'],
        ['Y2 (FY27/28)', '14', '0', '0', '0', '0', '14', '\u00a3280k'],
        ['Y3 (FY28/29)', '20', '4', '0', '0', '0', '24', '\u00a3720k'],
        ['Y4', '28', '10', '2', '0', '0', '40', '\u00a31.66M'],
        ['Y5', '35', '18', '5', '2', '0', '60', '\u00a33.39M'],
        ['Y6', '40', '28', '10', '4', '1', '83', '\u00a35.94M'],
        ['Y7', '45', '38', '16', '7', '3', '109', '\u00a39.51M'],
        ['Y8', '48', '48', '22', '12', '5', '135', '\u00a313.76M'],
        ['Y9', '50', '58', '28', '16', '7', '159', '\u00a318.26M'],
        ['Y10', '52', '68', '35', '20', '10', '185', '\u00a324.46M'],
    ], col_widths=[2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 2.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. FINANCIAL PROJECTIONS — 3-YEAR DETAILED
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '11. Financial Projections \u2014 3-Year Detailed', level=1)

    doc.add_paragraph(
        'All projections assume the consortium forms in winter 2026, with first contributions received '
        'in Q3 FY 26/27 (October\u2013December 2026). The 24-month build programme runs from winter 2026 '
        'to winter 2028, with production release and SaaS sales commencing in FY 28/29.'
    )

    add_heading_styled(doc, 'Revenue Forecast', level=2)
    add_styled_table(doc, ['Revenue Stream', 'FY 26/27', 'FY 27/28', 'FY 28/29'], [
        ['Consortium (10 x \u00a315k over 24m)', '\u00a375,000', '\u00a375,000', '\u2014'],
        ['SaaS \u2014 Small HAs', '\u2014', '\u00a350,000', '\u00a3400,000'],
        ['SaaS \u2014 Mid-market HAs', '\u2014', '\u2014', '\u00a3320,000'],
        ['Total Revenue', '\u00a375,000', '\u00a3125,000', '\u00a3720,000'],
    ], col_widths=[7, 3, 3, 3])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'Note: Years 1-2 are build years. No SaaS sales in Year 1. '
        '1-2 early SaaS customers in Year 2 (from consortium demo pipeline). '
        'Year 3 is the first real sales year once the product is production-ready.'
    )
    run.italic = True
    run.font.size = Pt(9)

    add_heading_styled(doc, 'Cost Forecast \u2014 People', level=2)
    add_styled_table(doc, ['Role', 'Location', 'FY 26/27', 'FY 27/28', 'FY 28/29'], [
        ['Founder / CTO', 'UK', '\u00a395,000', '\u00a395,000', '\u00a395,000'],
        ['Employer NI & pension (UK roles)', '', '\u00a36,000', '\u00a316,000', '\u00a326,000'],
        ['Full-Stack Developers (2\u21923\u21925)', 'Offshore', '\u00a350,000', '\u00a375,000', '\u00a3125,000'],
        ['VAT on offshore services (20%)', '', '\u00a310,000', '\u00a315,000', '\u00a325,000'],
        ['Sales / BD (from Y2)', 'UK', '\u2014', '\u00a345,000', '\u00a350,000'],
        ['DevOps Engineer (from Y3)', 'Offshore', '\u2014', '\u2014', '\u00a328,000'],
        ['Customer Success (from Y3)', 'UK part-time', '\u2014', '\u2014', '\u00a330,000'],
        ['Subtotal People', '', '\u00a3161,000', '\u00a3246,000', '\u00a3379,000'],
    ], col_widths=[5.5, 3, 2.5, 2.5, 2.5])
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'UK employer costs include 13.8% NI + 3% pension auto-enrolment. '
        'Offshore services via Yantra Works incur 20% VAT (reverse charge mechanism). '
        'Founder salary of \u00a395k structured as salary + dividends for tax efficiency.'
    )
    run.italic = True
    run.font.size = Pt(8)

    add_heading_styled(doc, 'Cost Forecast \u2014 Other', level=2)
    add_styled_table(doc, ['Category', 'FY 26/27', 'FY 27/28', 'FY 28/29'], [
        ['Independent QA/testing company', '\u00a315,000', '\u00a320,000', '\u00a325,000'],
        ['Cloud hosting & AI APIs', '\u00a314,000', '\u00a323,000', '\u00a350,000'],
        ['Legal, compliance, Cyber Essentials', '\u00a312,000', '\u00a318,000', '\u00a327,000'],
        ['Marketing & conferences', '\u00a310,000', '\u00a320,000', '\u00a345,000'],
        ['Grant funding applications', '\u00a315,000', '\u00a315,000', '\u00a315,000'],
        ['Admin & insurance', '\u00a38,000', '\u00a310,000', '\u00a316,000'],
        ['Subtotal Other', '\u00a374,000', '\u00a3106,000', '\u00a3178,000'],
    ], col_widths=[7, 3, 3, 3])

    add_heading_styled(doc, 'Profit & Loss Summary \u2014 3-Year', level=2)
    add_styled_table(doc, ['Line Item', 'FY 26/27', 'FY 27/28', 'FY 28/29'], [
        ['Total Revenue', '\u00a375,000', '\u00a3125,000', '\u00a3720,000'],
        ['Total Costs', '(\u00a3235,000)', '(\u00a3352,000)', '(\u00a3557,000)'],
        ['Net Profit / (Loss)', '(\u00a3160,000)', '(\u00a3227,000)', '+\u00a3163,000'],
        ['Cumulative Position', '(\u00a3160,000)', '(\u00a3387,000)', '(\u00a3224,000)'],
    ], col_widths=[6, 3.5, 3.5, 3.5])

    doc.add_paragraph()
    add_bold_para(doc,
        'The cumulative funding requirement is \u00a3387,000 at the end of Year 2, reducing to '
        '\u00a3224,000 by end of Year 3. Funded by: consortium contributions (\u00a3150,000) + '
        'investment (\u00a3250,000) = \u00a3400,000 total available. The gap closes further in Year 4 '
        'as mid-market sales accelerate. The business becomes self-sustaining from Year 4 onward, '
        'generating 40%+ gross margins.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 12. FINANCIAL PROJECTIONS — 10-YEAR GROWTH MODEL
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '12. Financial Projections \u2014 10-Year Growth Model', level=1)

    doc.add_paragraph(
        'The 10-year model shows how revenue scales dramatically as the platform moves into higher-value '
        'market segments, while costs grow more slowly due to platform leverage and AI-assisted operations.'
    )

    chart_path = os.path.join(CHART_DIR, 'info_arr_growth.png')
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
    else:
        chart_path = create_arr_growth_chart()
        doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_styled_table(doc, ['Year', 'Headcount', 'Total Costs', 'Revenue', 'Net Profit / (Loss)'], [
        ['Y1 (FY 26/27)', '5 + testing co', '\u00a3375k', '\u00a3200k', '(\u00a3175k)'],
        ['Y2 (FY 27/28)', '7', '\u00a3490k', '\u00a3280k', '(\u00a3210k)'],
        ['Y3 (FY 28/29)', '12', '\u00a3750k', '\u00a3720k', '(\u00a330k) ~breakeven'],
        ['Y4', '18', '\u00a31.1M', '\u00a31.66M', '+\u00a3560k'],
        ['Y5', '25', '\u00a31.8M', '\u00a33.39M', '+\u00a31.59M'],
        ['Y6', '35', '\u00a32.8M', '\u00a35.94M', '+\u00a33.14M'],
        ['Y7', '50', '\u00a34.2M', '\u00a39.51M', '+\u00a35.31M'],
        ['Y8', '65', '\u00a35.8M', '\u00a313.76M', '+\u00a37.96M'],
        ['Y9', '80', '\u00a37.2M', '\u00a318.26M', '+\u00a311.06M'],
        ['Y10', '95', '\u00a38.5M', '\u00a324.46M', '+\u00a315.96M'],
    ], col_widths=[3, 3, 3, 3, 4])

    doc.add_paragraph()
    chart_path = create_revenue_cost_10yr_chart()
    doc.add_picture(chart_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_bold_para(doc,
        'Cumulative profit through 10 years: approximately \u00a345M. '
        'Investment recovered in Year 4. By Year 10: 242 customers, \u00a332.4M ARR, 95 employees, '
        'and a platform serving the entire UK social housing market from small HAs to the G15.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 13. TEAM & RESOURCING
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '13. Team & Resourcing', level=1)

    doc.add_paragraph(
        'The team scales with revenue \u2014 starting lean and growing only as customer demand requires. '
        'This is a capital-efficient approach that avoids the common startup mistake of hiring ahead of revenue.'
    )

    add_heading_styled(doc, 'Year 1 Team (5 people + testing company)', level=2)
    add_styled_table(doc, ['Role', 'Location', 'Annual Cost', 'Notes'], [
        ['CTO / Architect', 'UK', '\u00a395,000', 'Could be founder at reduced salary'],
        ['Full-Stack Developer x3', 'Offshore (Yantra)', '\u00a390,000 (\u00a330k each)', 'Supervised by CTO'],
        ['Product Manager / BA (0.3 FTE)', 'UK contract', '\u00a325,000', 'Housing domain expertise'],
        ['Independent QA/Testing co', 'UK', '\u00a360,000', 'Functional, security, pen testing'],
        ['Year 1 Total', '', '\u00a3270,000', ''],
    ], col_widths=[5.5, 3, 3.5, 5])

    add_heading_styled(doc, 'Year 2 Team (7 people)', level=2)
    add_styled_table(doc, ['Addition', 'Location', 'Annual Cost', 'Notes'], [
        ['DevOps Engineer', 'Offshore', '\u00a328,000', 'Infrastructure, monitoring, deployment'],
        ['Customer Success', 'UK', '\u00a330,000', 'Onboarding, training, support'],
        ['Year 2 Total (all roles)', '', '~\u00a3330,000', 'Including Y1 team'],
    ], col_widths=[5.5, 3, 3.5, 5])

    add_heading_styled(doc, 'Year 3 Team (12 people)', level=2)
    add_styled_table(doc, ['Addition', 'Location', 'Annual Cost', 'Notes'], [
        ['Additional developers x3', 'Offshore', '\u00a390,000 (\u00a330k each)', 'Scale for mid-market features'],
        ['Sales Lead', 'UK', '\u00a365,000', 'Dedicated sales for SaaS growth'],
        ['Data Engineer', 'Offshore', '\u00a330,000', 'Data migration, reporting, analytics'],
        ['Year 3 Total (all roles)', '', '~\u00a3475,000', 'Full team with increments'],
    ], col_widths=[5.5, 3, 3.5, 5])

    add_heading_styled(doc, 'Scaling to 95 People by Year 10', level=2)
    doc.add_paragraph('The team grows proportionally with customer base and revenue:')
    add_styled_table(doc, ['Phase', 'Headcount', 'Key Additions'], [
        ['Years 1\u20132 (Prove)', '5\u20137', 'Core engineering team + testing company'],
        ['Years 3\u20134 (Grow)', '12\u201318', 'Sales, data engineering, additional devs, customer success'],
        ['Years 5\u20136 (Scale)', '25\u201335', 'Enterprise sales, implementation team, dedicated security'],
        ['Years 7\u20138 (Expand)', '50\u201365', 'Product teams per segment, professional services, account mgmt'],
        ['Years 9\u201310 (Dominate)', '80\u201395', 'G15 dedicated team, international R&D, exec leadership'],
    ], col_widths=[4, 2.5, 10.5])

    add_heading_styled(doc, 'Yantra Works Resourcing Capability', level=2)
    for point in [
        'UK-registered company with onshore account management and 20+ years\' senior experience',
        'Offshore teams available with 1 week notice; option to bring onshore (12-month minimum)',
        'Cost advantage: offshore senior developer at \u00a325k\u2013\u00a335k/yr vs \u00a370k\u2013\u00a387k UK equivalent (60\u201370% saving)',
        'All work supervised by experienced UK-based engineers and independent testing company',
    ]:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 14. MINIMUM VIABLE BUDGET SCENARIO
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '14. Minimum Viable Budget Scenario', level=1)

    doc.add_paragraph(
        'What could be achieved with consortium revenue alone (\u00a3150,000 over 24 months) and no external investment?'
    )

    add_styled_table(doc, ['Item', 'Cost (24 months)'], [
        ['Architect (part-time / founder sweat equity)', '\u00a30'],
        ['2 x offshore developers (Yantra)', '\u00a360,000'],
        ['Testing company (reduced scope)', '\u00a325,000'],
        ['Cloud hosting & AI (Claude API)', '\u00a315,000'],
        ['Legal, compliance, marketing', '\u00a320,000'],
        ['Total', '~\u00a3120,000'],
    ], col_widths=[10, 5])

    doc.add_paragraph()
    doc.add_paragraph('This scenario is viable but very constrained:')
    for constraint in [
        'Timeline extends to 3\u20134 years (vs 2 years with investment)',
        'Limited feature set \u2014 only core modules built to production standard',
        'No independent testing company for most of the build (security risk)',
        'No sales and marketing budget \u2014 growth relies entirely on word of mouth',
        'Founder takes no salary for 2+ years (significant personal risk)',
        'No ISO 27001 certification (limits procurement eligibility for mid-market and above)',
        'Cannot enter mid-market or enterprise segments \u2014 stuck in small HA segment only',
        'Single point of failure on the architect/founder',
    ]:
        doc.add_paragraph(constraint, style='List Bullet')

    doc.add_paragraph()
    add_bold_para(doc,
        'Conclusion: The minimum viable budget produces a working product for small HAs eventually, '
        'but cannot access the mid-market, large, or G15 segments. The \u00a3250,000 investment '
        'accelerates time-to-market by 12\u201318 months AND unlocks the path to \u00a332.4M ARR \u2014 '
        'the difference between a small niche product and a sector-defining platform.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 15. RISK ANALYSIS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '15. Risk Analysis', level=1)

    add_styled_table(doc, ['Risk', 'Likelihood', 'Impact', 'Mitigation'], [
        ['Consortium fails to reach 10 members', 'Medium', 'High',
         'Viable at 7 members (\u00a3105k). Low \u00a315k ask reduces barrier. 3C network provides pipeline. Grant funding pursued in parallel.'],
        ['Slower customer acquisition post-consortium', 'Medium', 'Medium',
         'Conservative Year 3 target (24 = 1.6% of TAM). NEC OHMS end-of-life creating forced migration wave.'],
        ['Mid-market entry takes longer than projected', 'Medium', 'Medium',
         'Small HA revenue covers costs through Year 3. Mid-market is additive, not required for survival.'],
        ['G15 contracts harder to win than projected', 'Medium', 'Low',
         'G15 revenue is Year 6+ upside. The model is profitable from Year 4 without any G15 contracts.'],
        ['Data migration complexity', 'High', 'Medium',
         'HACT v3.5 standard reduces variability. 3-month parallel running. Migration scripts for top 4 source systems.'],
        ['Claude API costs exceed estimates', 'Low', 'Low',
         'Haiku handles 80% of queries at 1/3 cost. Total AI cost is ~\u00a3256/month even at 242 customers.'],
        ['Incumbent vendor response', 'Medium', 'Low',
         'Legacy vendors cannot easily retrofit AI into 30-year-old codebases. Their strategy suggests buy-not-build.'],
        ['Security breach', 'Low', 'High',
         'Independent testing company, Cyber Essentials+, CREST pen testing, UK-only data, ISO 27001 from Y2.'],
        ['Key person dependency', 'Medium', 'Medium',
         'Code documented, 625 tests, GitHub access for consortium. Team scales to 95 by Y10.'],
        ['Funding shortfall', 'Low', 'High',
         'Minimum viable budget scenario shows viability at \u00a3150k. Multiple paths: grants, investment, revenue.'],
    ], col_widths=[4, 2, 2, 9])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 16. COMPARABLE COMPANIES & EXIT SCENARIOS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '16. Comparable Companies & Exit Scenarios', level=1)

    add_heading_styled(doc, 'UK Housing PropTech Benchmarks', level=2)
    add_styled_table(doc, ['Company', 'Focus', 'Funding Raised', 'Revenue', 'Customers', 'Valuation'], [
        ['Plentific', 'Repairs & maintenance SaaS', '~$136\u2013166M (incl. $100M Series C)', '$21.1M', '100', 'Not disclosed'],
        ['Switchee', 'Smart thermostat / IoT', '$20.1M (6 rounds)', '$11.6M', '130+', '$1.63B (Jul 2024)'],
        ['FaultFixers', 'Maintenance management', '\u00a3421K (grant + investment)', 'N/A', 'Growing', 'Early stage'],
        ['Localz', 'Last-mile engagement', '$7.44M (3 rounds)', 'N/A', '10+', 'Acquired for $6.2M'],
    ], col_widths=[2.5, 3.5, 4.5, 2.5, 2, 3])
    add_source_note(doc, 'Sources: TechCrunch, GetLatka, UK Tech News, Descartes Systems Group')

    doc.add_paragraph()
    doc.add_paragraph(
        'Switchee\'s $1.63 billion valuation from $20M in funding demonstrates the extraordinary multiples '
        'achievable in UK housing PropTech. Plentific\'s $100M Series C shows deep investor appetite '
        'for the sector. SocialHomes.Ai targets the core HMS market \u2014 a larger opportunity than either '
        'Switchee (IoT) or Plentific (repairs only).'
    )

    add_heading_styled(doc, 'Exit Scenarios', level=2)
    add_styled_table(doc, ['Scenario', 'Timeline', 'Indicative Value', 'Notes'], [
        ['A. Organic growth (base case)', 'Ongoing', '\u00a332.4M ARR by Year 10', 'Self-sustaining from Year 4, highly profitable'],
        ['B. Strategic acquisition', 'Year 5\u20137', '\u00a320M\u2013\u00a376M (5\u20138x ARR)', 'MRI, Aareon, Capita, or PE firm'],
        ['C. Full-scale exit at Year 10', 'Year 10', '~\u00a3260M (8x ARR)', '242 customers, market leader'],
        ['D. Series A raise', 'Year 3\u20134', '\u00a35M\u2013\u00a310M at \u00a320M\u2013\u00a340M pre-money', 'Accelerate mid-market entry'],
        ['E. Grant + organic', 'Year 1\u20133', 'Reduced dilution', 'Innovate UK / PropTech Fund'],
    ], col_widths=[4, 2.5, 5, 5.5])

    doc.add_paragraph()
    add_bold_para(doc,
        'The maths: \u00a3250k invested. Year 10 ARR: \u00a332.4M. At 8x revenue multiple: \u00a3260M valuation. '
        'That is a 1,000x return. Even at a conservative 5x multiple, the valuation is \u00a3122M \u2014 '
        'a 254x return. These multiples are standard for high-growth vertical SaaS businesses '
        'in regulated sectors.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 17. INVESTMENT ASK
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '17. Investment Ask', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Yantra Works Ltd is seeking \u00a3250,000 in investment.')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL

    doc.add_paragraph()
    doc.add_paragraph(
        'This investment, combined with \u00a3150,000 in consortium contributions, funds the bridge from '
        'a 50\u201355% complete prototype to a production platform serving the entire UK social housing market. '
        'The total funding available is \u00a3630,000, with breakeven achieved in Year 3 and strong '
        'profitability from Year 4 onwards.'
    )

    add_heading_styled(doc, 'Detailed Use of Funds', level=2)
    add_styled_table(doc, ['Category', 'Amount', '% of Total', 'What It Covers'], [
        ['Production engineering', '\u00a3240,000', '50%', '3 fullstack developers + CTO oversight for 24 months'],
        ['Independent QA/testing', '\u00a360,000', '12.5%', 'Functional, security, accessibility, CREST pen testing'],
        ['Cloud infrastructure & AI', '\u00a348,000', '10%', 'GCP hosting, Claude API, monitoring, dev tools'],
        ['Sales & marketing', '\u00a350,000', '10.4%', 'Conferences, content, 3C partnership, early sales'],
        ['Legal, compliance, ISO 27001', '\u00a342,000', '8.8%', 'Contracts, GDPR DPIAs, Cyber Essentials+, ISO prep'],
        ['Contingency', '\u00a340,000', '8.3%', 'Buffer for delays, scope changes, market conditions'],
        ['Total', '\u00a3250,000', '100%', ''],
    ], col_widths=[4.5, 2.5, 2, 8])

    add_heading_styled(doc, 'Investment Structure Options', level=2)
    add_styled_table(doc, ['Option', 'Structure', 'Suited To'], [
        ['A. Revenue-based financing', '\u00a3250k loan, repaid at 8\u201312% of monthly revenue until 1.5x returned', 'Banks / credit facilities'],
        ['B. Convertible loan note', '\u00a3250k at agreed interest, convertible to equity at next raise', 'Angel investors'],
        ['C. Equity investment', '\u00a3250k for negotiable minority stake in Yantra Works Ltd', 'Strategic investors'],
        ['D. Blended (grant + investment)', 'Grant + smaller private investment = lowest dilution', 'See grant options below'],
    ], col_widths=[4, 7.5, 5.5])

    add_heading_styled(doc, 'Grant Funding Options', level=2)
    add_styled_table(doc, ['Grant Programme', 'Available', 'Fit', 'Notes'], [
        ['Innovate UK Smart Grants', 'Up to \u00a3500k', 'Strong', 'Disruptive innovation in regulated sector'],
        ['PropTech Innovation Fund', 'Variable', 'Strong', 'Directly targeted at property technology'],
        ['Social Housing Decarbonisation Fund', 'Variable', 'Tangential', 'Energy/damp features may qualify'],
        ['UKRI / AHRC', 'Variable', 'Moderate', 'AI in public services research funding'],
    ], col_widths=[5, 2.5, 2, 7.5])

    add_heading_styled(doc, 'Return Profile', level=2)
    add_styled_table(doc, ['Metric', 'Value'], [
        ['Investment amount', '\u00a3250,000'],
        ['Consortium contributions', '\u00a3150,000 (10 x \u00a315,000 over 24 months)'],
        ['Total funding', '\u00a3630,000'],
        ['Breakeven year', 'Year 3 (FY 28/29)'],
        ['Investment recovered', 'Year 4'],
        ['Year 5 ARR', '\u00a33.39M'],
        ['Year 10 ARR', '\u00a324.46M'],
        ['Year 10 valuation (8x ARR)', '\u00a3260M'],
        ['Return multiple on \u00a3250k', '~1,000x'],
        ['Cumulative profit (10 years)', '~\u00a345M'],
        ['All IP remains with', 'Yantra Works Ltd'],
    ], col_widths=[6, 10])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 18. KEY MILESTONES
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, '18. Key Milestones', level=1)

    add_styled_table(doc, ['Date', 'Milestone', 'Measure of Success'], [
        ['Q2 2026', 'Investment secured / grant application', 'Funds available'],
        ['Q2\u2013Q3 2026', 'Sales pipeline built', '30+ qualified leads from conferences / 3C'],
        ['Winter 2026', 'Consortium Board formed', '10 housing associations committed'],
        ['Q4 2026', 'Production team hired', '5 people in place (CTO + 3 devs + PM)'],
        ['Q1 2027', 'Full CRUD operations complete', 'All modules support create, update, delete'],
        ['Q2 2027', 'Alpha release to 3 consortium members', 'Parallel running with existing HMS'],
        ['Q3 2027', 'Cyber Essentials+ certification', 'Certification obtained'],
        ['Q4 2027', 'Beta release to all 10 consortium', 'UAT underway'],
        ['Q1 2028', 'Multi-tenancy complete', '10 isolated tenant environments'],
        ['Q1 2028', 'ISO 27001 implementation begins', 'Gap analysis complete'],
        ['Q2 2028', 'First non-consortium SaaS customer', 'Contract signed'],
        ['Q3 2028', 'Production v1.0 release', 'All consortium members live'],
        ['Q1 2029', 'G-Cloud listing achieved', 'Listed on Digital Marketplace'],
        ['Mid 2029', 'ISO 27001 certification', 'Independent audit passed'],
        ['Q4 2029', '24 customers live, ~breakeven', 'Year 3 targets met'],
        ['2030', 'First mid-market customer (5k+ homes)', 'Contract at \u00a380k/yr tier'],
        ['2031', 'First enterprise customer (30k+ homes)', 'Contract at \u00a3250k/yr tier'],
        ['2032', 'First G15 member signed', '\u00a3500k/yr contract \u2014 sector validation'],
        ['2033\u201336', 'Market leadership established', '242 customers, \u00a332.4M ARR'],
    ], col_widths=[2.5, 6, 8.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX A: ABOUT YANTRA WORKS
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, 'Appendix A: About Yantra Works Ltd', level=1)

    doc.add_paragraph(
        'Yantra Works Ltd (yantra.works) is a UK-registered technology consultancy providing:'
    )
    for service in [
        'Independent, vendor-agnostic advisory to CIOs and C-suite leaders',
        'Strategy-to-prototype rapid validation \u2014 working with senior leadership to rapidly prototype solutions',
        'Vendor selection and enablement \u2014 choosing the right technologies and governing suppliers',
        'Flexible resourcing (onshore, offshore, hybrid) with 20+ years\' senior experience',
    ]:
        doc.add_paragraph(service, style='List Bullet')

    add_heading_styled(doc, 'The Yantra Framework: Agentic AI Development', level=2)
    doc.add_paragraph(
        'Yantra Works has developed a proprietary framework for agentic AI software development. '
        'Seven specialised AI agents work in a structured pipeline under human supervision:'
    )
    for agent in [
        'Business Analyst agent: Creates specifications and execution plans from requirements',
        'Fullstack Developer agent: Implements backend and frontend code following architecture patterns',
        'Frontend Developer agent: Builds responsive UI components with Tailwind CSS and React',
        'UI/UX Designer agent: Audits designs for consistency, accessibility, and usability',
        'QA Tester agent: Writes and runs automated tests (Playwright, Vitest)',
        'Cyber Security agent: Audits code for OWASP Top 10, secret scanning, dependency vulnerabilities',
        'DevOps agent: Configures CI/CD pipelines, Docker, Cloud Run deployment',
    ]:
        doc.add_paragraph(agent, style='List Bullet')

    doc.add_paragraph(
        'This framework reduces development costs by approximately 40% compared to traditional approaches, '
        'while an independent testing company provides human verification.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Key personnel:')
    doc.add_paragraph('Rajiv Peter \u2014 Founder & Managing Director. Contact: rajiv@yantra.works', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Partners:')
    doc.add_paragraph(
        '3C \u2014 Sector engagement partner with deep housing association relationships '
        'and implementation expertise. Responsible for consortium recruitment and housing domain advisory.',
        style='List Bullet',
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX B: GLOSSARY
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, 'Appendix B: Glossary for Non-Housing Readers', level=1)

    glossary = [
        ('Housing Association (HA)', 'Not-for-profit organisation that provides social housing'),
        ('Registered Provider (RP)', 'Formal regulatory term for a housing association registered with the RSH'),
        ('HMS', 'Housing Management System \u2014 the core operational software used by housing providers'),
        ('RSH', 'Regulator of Social Housing \u2014 the government regulator for social housing in England'),
        ('TSM', 'Tenant Satisfaction Measures \u2014 22 mandatory satisfaction metrics collected annually'),
        ('Awaab\'s Law', 'Legal requirement for timed response to damp, mould, and hazard reports'),
        ('Big 6', 'Six key safety compliance areas: Gas, Electrical, Fire, Asbestos, Legionella, Lifts'),
        ('G15', 'London\'s 15 largest housing associations, collectively managing 800,000+ homes'),
        ('HACT', 'Housing Associations\' Charitable Trust \u2014 sets the UK Housing Data Standard'),
        ('SOR Code', 'Schedule of Rates \u2014 standardised pricing for repair work'),
        ('Void', 'An empty property between tenancies'),
        ('Arrears', 'Unpaid rent owed by a tenant'),
        ('Universal Credit (UC)', 'UK welfare payment system that replaced Housing Benefit'),
        ('ALMO', 'Arms-Length Management Organisation \u2014 council housing managed at arm\'s length'),
        ('PropTech', 'Property Technology \u2014 technology companies serving the real estate sector'),
        ('ARR', 'Annual Recurring Revenue \u2014 predictable yearly revenue from subscriptions'),
        ('SaaS', 'Software as a Service \u2014 cloud-hosted subscription software'),
        ('G-Cloud', 'UK government procurement framework for cloud services'),
        ('GDPR', 'General Data Protection Regulation \u2014 EU/UK data protection law'),
        ('DPIA', 'Data Protection Impact Assessment \u2014 required for processing sensitive personal data'),
        ('ISO 27001', 'International standard for information security management systems'),
        ('Claude', 'Anthropic\'s large language model (LLM) used as the AI engine in SocialHomes.Ai'),
        ('Agentic AI', 'AI systems that can take autonomous actions within defined boundaries under supervision'),
        ('Multi-tenancy', 'Architecture where one application serves multiple organisations with isolated data'),
        ('CREST', 'Council of Registered Ethical Security Testers \u2014 UK penetration testing standard'),
    ]
    add_styled_table(doc, ['Term', 'Meaning'], glossary, col_widths=[5, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # APPENDIX C: SOURCES
    # ══════════════════════════════════════════════════════════════

    add_heading_styled(doc, 'Appendix C: Sources & References', level=1)

    sources = [
        ('Market Data', [
            'RSH Statistical Data Return 2024-25 (GOV.UK)',
            'RSH Global Accounts of Private Registered Providers 2025 (GOV.UK)',
            'RSH Value for Money Metrics and Reporting 2025 (GOV.UK)',
            'ARCH Housing \u2014 Association of Retained Council Housing',
            'HACT UK Housing Data Standards v3.5',
            'Housing Today \u2014 Largest 50 Housing Associations 2025',
            'G15 member annual reports and accounts 2024-25',
        ]),
        ('Regulatory', [
            'GOV.UK \u2014 Awaab\'s Law Guidance for Social Landlords',
            'Social Housing (Regulation) Act 2023',
            'RSH Tenant Satisfaction Measures 2024-25 Headline Report',
            'Housing Ombudsman Complaint Handling Code 2024',
            'Housing Ombudsman Annual Complaints Review 2024-25',
            'RSH Consumer Regulation Review 2023-24',
        ]),
        ('Procurement & Pricing', [
            'Sector benchmarks for HMS costs across all provider sizes',
            'Find a Tender Service \u2014 HMS contract notices 2024-25',
            'Contracts Finder \u2014 Housing management software awards',
            'CCS Framework RM6259 \u2014 Vertical Application Solutions',
        ]),
        ('AI & Technology', [
            'Anthropic Claude API Pricing (anthropic.com)',
            'Anthropic Claude Model Cards \u2014 Sonnet 4, Haiku 4.5',
            'Google Cloud Run Pricing (cloud.google.com)',
            'Google Cloud Firestore Pricing (cloud.google.com)',
        ]),
        ('Market Analysis', [
            'Market Research Future \u2014 UK PropTech Market Report 2024',
            'Grand View Research \u2014 UK PropTech Market Size & Trends',
            'Beauhurst \u2014 Top 10 PropTech Companies in the UK',
            'CRETI \u2014 2024 PropTech Venture Capital Report',
        ]),
        ('Comparable Companies', [
            'TechCrunch \u2014 Plentific $100M Series C',
            'UK Tech News \u2014 Switchee $1.63B valuation (July 2024)',
            'GetLatka \u2014 Plentific and Switchee revenue data',
            'Descartes Systems Group \u2014 Localz acquisition ($6.2M)',
        ]),
        ('Grant & Funding Programmes', [
            'Innovate UK Smart Grants (apply.innovateuk.org)',
            'PropTech Innovation Fund',
            'Social Housing Decarbonisation Fund (GOV.UK)',
            'UKRI / AHRC AI research funding programmes',
        ]),
        ('Standards & Certification', [
            'ISO/IEC 27001:2022 \u2014 Information Security Management Systems',
            'IASME \u2014 Cyber Essentials+ Certification',
            'OWASP Top 10 \u2014 Web Application Security Risks',
            'CREST \u2014 Council of Registered Ethical Security Testers',
        ]),
    ]

    for category, items in sources:
        doc.add_heading(category, level=3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # ── Footer ────────────────────────────────────────────────────

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared by Yantra Works Ltd | yantra.works | rajiv@yantra.works')
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('March 2026 | Version 3.0 | COMMERCIALLY SENSITIVE & CONFIDENTIAL')
    run.font.size = Pt(9)
    run.font.color.rgb = GREY

    # ── Save ──────────────────────────────────────────────────────

    output_path = os.path.join(OUTPUT_DIR, 'SocialHomesAi-Investment-Business-Plan-v4b-2026.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    build_document()
