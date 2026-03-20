#!/usr/bin/env python3
"""
Generate a 2-page investor brochure for SocialHomes.Ai.
Dense, punchy, no fluff — designed to be handed to investors.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

TEAL = RGBColor(0, 170, 164)
DARK = RGBColor(30, 30, 30)
GREY = RGBColor(100, 100, 100)
WHITE = RGBColor(255, 255, 255)
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(OUTPUT_DIR, 'charts', 'yantra-logo.png')

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None, header_color='00AAA4', font_size=8):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(font_size)
        set_cell_shading(cell, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
                    if ri == len(rows) - 1:  # Last row bold
                        r.bold = True
            if ri % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)
    style.font.color.rgb = DARK

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ══════════════════════════════════════════════════════
    # PAGE 1
    # ══════════════════════════════════════════════════════

    # Header with logo
    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Inches(0.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    run = p.add_run('SocialHomes.Ai')
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL
    run.bold = True
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('AI-Native Housing Management for UK Social Housing')
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('Investor Summary  |  March 2026  |  Commercially Sensitive')
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.italic = True
    p.paragraph_format.space_after = Pt(8)

    # ── WHY THIS PRODUCT? ──
    h = doc.add_heading('Why SocialHomes.Ai?', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(12)

    doc.add_paragraph(
        'UK social housing manages 4.5 million homes for 9 million people. Four new laws since 2023 '
        '(Awaab\u2019s Law, Tenant Satisfaction Measures, Complaint Handling Code, RSH Consumer Standards) '
        'demand better data, faster response times, and real accountability. '
        'Yet 80% of the sector\u2019s technology was designed in the 1990s.'
    ).paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run('Smaller housing associations \u2014 the ones closest to their communities \u2014 '
                     'pay \u00a350k\u2013\u00a3100k/year for software that wasn\u2019t built for them. ')
    run.font.size = Pt(9)
    run = p.add_run('SocialHomes.Ai is the first AI-native, cloud-native housing management system '
                     'designed specifically for this underserved segment, at 50\u201370% lower cost.')
    run.bold = True
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)

    # ── WHAT'S DIFFERENT ──
    h = doc.add_heading('What Makes This Different', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    diffs = [
        ('AI-native, not AI-bolted-on', 'Intelligence on every screen \u2014 predicts arrears, identifies vulnerable tenants, flags compliance risks before they become crises'),
        ('Self-service onboarding', 'No other HMS vendor offers this. Sign up \u2192 import data \u2192 go live in weeks, not the industry-standard 6\u201318 months'),
        ('Combines public + organisational data', 'Auto-pulls EPC, weather, crime, flood, deprivation data and enriches every property and tenant record'),
        ('50\u201370% cheaper', '\u00a320k/year for a 2,500-home HA vs \u00a355k\u2013\u00a366k for Civica/Aareon/NEC (G-Cloud published pricing)'),
        ('Consortium-governed', '10 founding HAs co-invest \u00a315k each and shape the roadmap via a Board seat'),
    ]
    for title, desc in diffs:
        p = doc.add_paragraph()
        run = p.add_run(f'\u2022 {title}. ')
        run.bold = True
        run.font.size = Pt(8.5)
        run = p.add_run(desc)
        run.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    # ── MARKET ──
    h = doc.add_heading('Market', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    add_table(doc, ['', ''], [
        ['Total providers', '1,581 (managing 4.5M homes, \u00a327.4bn turnover)'],
        ['Target segment', '~400 providers with 1,000\u201330,000 homes'],
        ['Incumbent HMS cost', '\u00a350k\u2013\u00a3100k/yr for small HAs (G-Cloud data)'],
        ['SocialHomes.Ai price', '\u00a320k/yr (small) to \u00a3500k/yr (G15/mega)'],
        ['Key regulatory driver', 'Awaab\u2019s Law \u2014 unlimited fines for non-compliance'],
    ], col_widths=[4, 13], font_size=8)

    doc.add_paragraph()

    # ── COMPETITIVE LANDSCAPE ──
    h = doc.add_heading('Competition \u2014 Why We Win', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    add_table(doc, ['', 'Civica/NEC/MRI', 'Aareon QL', 'SocialHomes.Ai'], [
        ['Architecture', '1990s client-server', 'Hosted', 'Cloud-native (GCP)'],
        ['AI', 'Bolt-on module (\u00a3\u00a3\u00a3)', 'Limited', 'Native (every screen)'],
        ['Awaab\u2019s Law', 'Patch', 'Patch', 'Built-in from day 1'],
        ['Self-service setup', 'No (6\u201318 months)', 'No', 'Yes (weeks)'],
        ['Data standard', 'Proprietary', 'Proprietary', 'HACT v3.5 native'],
        ['Cost (2,500 homes)', '\u00a355\u2013\u00a366k/yr', '\u00a360k/yr', '\u00a320k/yr'],
    ], col_widths=[3.5, 4.5, 4, 5], font_size=8)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # PAGE 2
    # ══════════════════════════════════════════════════════

    h = doc.add_heading('10-Year Financial Projection', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(12)

    add_table(doc,
        ['Year', 'Customers', 'Revenue', 'Costs', 'Net', 'Valuation (8\u00d7)'],
        [
            ['Y1', '10', '\u00a375k', '\u00a3235k', '(\u00a3160k)', '\u00a30.6M'],
            ['Y2', '15', '\u00a3125k', '\u00a3352k', '(\u00a3227k)', '\u00a31M'],
            ['Y3', '27', '\u00a3720k', '\u00a3557k', '+\u00a3163k', '\u00a35.8M'],
            ['Y4', '48', '\u00a32.2M', '\u00a31.3M', '+\u00a3880k', '\u00a318M'],
            ['Y5', '77', '\u00a35M', '\u00a33M', '+\u00a32M', '\u00a340M'],
            ['Y6', '112', '\u00a312.5M', '\u00a37.5M', '+\u00a35M', '\u00a3100M'],
            ['Y7', '147', '\u00a315.8M', '\u00a39.5M', '+\u00a36.3M', '\u00a3126M'],
            ['Y8', '182', '\u00a321.4M', '\u00a312.9M', '+\u00a38.6M', '\u00a3172M'],
            ['Y9', '212', '\u00a326.8M', '\u00a316.1M', '+\u00a310.7M', '\u00a3214M'],
            ['Y10', '242', '\u00a332.4M', '\u00a319.5M', '+\u00a313M', '\u00a3260M'],
        ],
        col_widths=[1.5, 2, 2.5, 2.5, 2.5, 3], font_size=8)

    doc.add_paragraph()

    # ── RETURN PROFILE ──
    h = doc.add_heading('Return Profile', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    add_table(doc, ['Metric', 'Value'], [
        ['Investment sought', '\u00a3250,000'],
        ['Consortium co-investment', '\u00a3150,000 (10 HAs \u00d7 \u00a315k)'],
        ['Total funding', '\u00a3400,000'],
        ['Breakeven', 'Year 3 (\u00a3720k revenue)'],
        ['Year 6 ARR / valuation', '\u00a312.5M / \u00a3100M'],
        ['Year 10 ARR / valuation', '\u00a332.4M / \u00a3260M'],
        ['Return on \u00a3250k investment', '~1,000\u00d7 at Year 10 (at 8\u00d7 ARR)'],
        ['Gross margin from Year 4', '40%+'],
        ['Cumulative profit (10 years)', '~\u00a345M'],
    ], col_widths=[6, 11], font_size=8)

    doc.add_paragraph()

    # ── COMPARABLE SOFTWARE VALUATIONS ──
    h = doc.add_heading('Comparable UK Housing Tech Valuations', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    add_table(doc, ['Company', 'Focus', 'Funding', 'Revenue', 'Valuation', 'Customers'], [
        ['Switchee', 'IoT / smart thermostats', '\u00a320M', '\u00a311.6M', '\u00a31.63 BILLION', '130+'],
        ['Plentific', 'Repairs SaaS', '\u00a3136M', '\u00a321M', 'Not disclosed', '100'],
        ['Localz', 'Last-mile engagement', '\u00a37.4M', 'N/A', 'Acquired \u00a36.2M', '10+'],
        ['FaultFixers', 'Maintenance mgmt', '\u00a3421k', 'N/A', 'Early stage', 'Growing'],
        ['SocialHomes.Ai', 'Full HMS + AI', '\u00a3250k (target)', '\u00a375k (Y1)', '\u00a3100M (Y6 target)', '10 (Y1)'],
    ], col_widths=[2.5, 3, 2.5, 2, 3.5, 2], font_size=7.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'Switchee achieved a \u00a31.63 billion valuation from just \u00a320M in funding and 130 customers. '
        'SocialHomes.Ai targets a wider market (full HMS vs point solution) at a fraction of the capital.'
    )
    run.italic = True
    run.font.size = Pt(8)

    # ── SOCIAL IMPACT ──
    h = doc.add_heading('Social Impact', level=2)
    for r in h.runs:
        r.font.color.rgb = TEAL
        r.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(
        'Better housing systems save lives. ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(
        'When Awaab Ishak\u2019s mould complaints were tracked by a system with AI-powered countdown timers '
        'and automatic escalation, he would still be alive. When AI identifies that a vulnerable tenant '
        'hasn\u2019t been contacted in 6 months, a welfare check happens before a crisis. '
        'When predictive maintenance catches a failing boiler before winter, a family stays warm. '
        'Every pound saved on overpriced legacy software goes back into housing and services for '
        'the 9 million people who depend on social housing.'
    )
    run.font.size = Pt(9)

    doc.add_paragraph()

    # ── CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Rajiv Peter  |  Founder & CTO  |  rajiv@yantra.works  |  yantra.works')
    run.font.size = Pt(9)
    run.font.color.rgb = TEAL
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Yantra Works Ltd  |  Commercially Sensitive & Confidential')
    run.font.size = Pt(7)
    run.font.color.rgb = GREY

    # Save
    path = os.path.join(OUTPUT_DIR, 'SocialHomesAi-Investor-Brochure-2026.docx')
    doc.save(path)
    print(f'Brochure saved to: {path}')

if __name__ == '__main__':
    build()
